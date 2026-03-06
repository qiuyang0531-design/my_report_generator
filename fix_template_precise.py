# -*- coding: utf-8 -*-
"""
精确修复 template.docx - 将 Jinja2 标签紧凑化
策略：将独立的标签段落合并到其相邻的内容段落中
"""
from docx import Document
import shutil

def fix_template_precise():
    """精确修复模板布局"""
    # 备份
    shutil.copy('template.docx', 'template_backup_before_fix2.docx')
    print("已备份template.docx")

    doc = Document('template.docx')

    print('=== 开始精确修复 ===')

    # 策略：遍历所有段落，将独立的标签合并到相邻段落
    i = 0
    while i < len(doc.paragraphs):
        para = doc.paragraphs[i]
        text = para.text.strip()

        # 检查是否是独立的标签段落（只有Jinja2标签）
        is_standalone_tag = (
            text.startswith('{%') and
            text.endswith('%}') and
            not ('}}' in text or '{{' in text)  # 不包含变量输出
        )

        if is_standalone_tag:
            # 检查前一段和后一段
            prev_para = doc.paragraphs[i - 1] if i > 0 else None
            next_para = doc.paragraphs[i + 1] if i + 1 < len(doc.paragraphs) else None

            # 决定合并方向
            merged = False

            # 如果是 {% set %} 或 {% if %}，尝试合并到前一段
            if any(kw in text for kw in ['{% set', '{% if']) and prev_para:
                # 将标签添加到前一段的末尾
                if prev_para.runs:
                    prev_para.runs[-1].text = prev_para.runs[-1].text + text
                    print(f"合并到前一段: {text[:50]}...")
                    # 清空当前段落
                    for run in para.runs:
                        run.text = ''
                    merged = True

            # 如果是 {% endfor %} 或 {% endif %}，尝试合并到后一段或前一段
            elif any(kw in text for kw in ['{% endfor', '{% endif']):
                # 优先合并到后一段的开头
                if next_para and next_para.runs:
                    next_para.runs[0].text = text + next_para.runs[0].text
                    print(f"合并到后一段: {text[:50]}...")
                    for run in para.runs:
                        run.text = ''
                    merged = True
                elif prev_para and prev_para.runs:
                    prev_para.runs[-1].text = prev_para.runs[-1].text + text
                    print(f"合并到前一段: {text[:50]}...")
                    for run in para.runs:
                        run.text = ''
                    merged = True

            # 如果是 {% for %}，尝试合并到后一段（循环内容）
            elif '{% for' in text and next_para:
                if next_para.runs:
                    next_para.runs[0].text = text + next_para.runs[0].text
                    print(f"合并for标签到后一段: {text[:50]}...")
                    for run in para.runs:
                        run.text = ''
                    merged = True

        i += 1

    print("\n=== 检查量化方法说明部分 ===")
    for i, para in enumerate(doc.paragraphs[185:220], start=185):
        text = para.text.strip()
        if text or not text:
            status = '[EMPTY]' if not text else '[CONTENT]'
            preview = text[:80] if text else ''
            if text:
                print(f'{i}: {preview}')

    # 保存
    doc.save('template_fixed2.docx')
    print("\n已保存到 template_fixed2.docx")

    # 替换原文件
    import os
    if os.path.exists('template.docx'):
        os.remove('template.docx')
    shutil.move('template_fixed2.docx', 'template.docx')
    print("已替换 template.docx")

if __name__ == '__main__':
    try:
        fix_template_precise()
        print("\n=== 精确修复完成 ===")
    except Exception as e:
        print(f"错误: {e}")
        import traceback
        traceback.print_exc()
