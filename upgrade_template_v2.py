# -*- coding: utf-8 -*-
"""
升级 template.docx 以适配新的 report_config.py 嵌套结构 v2
正确处理 Word 段落结构
"""
import sys
from docx import Document
from docx.shared import Pt
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


def upgrade_template_v2():
    """
    升级模板，正确处理 Word 段落结构
    """
    print("开始升级 template.docx v2...")

    doc = Document('template.docx')

    # 我们需要找到特定的段落并替换整个块
    # 由于 Word 文档的复杂性，我们采用以下策略：
    # 1. 找到量化方法说明章节的起始段落
    # 2. 删除旧的内容段落
    # 3. 添加新的结构化段落

    paragraphs_to_process = []
    i = 0
    while i < len(doc.paragraphs):
        para = doc.paragraphs[i]
        text = para.text.strip()

        # 找到量化方法说明章节
        if text == '量化方法说明':
            print(f"找到量化方法说明章节，位置: {i}")
            paragraphs_to_process.append(('start', i))

        # 找到需要替换的循环段落
        if '{% for source_name, method in quantification_methods.scope_1.items() %}' in text:
            print(f"找到范围一循环，位置: {i}")
            paragraphs_to_process.append(('scope1_loop', i))

        if '{% for source_name, method in quantification_methods.scope_2.items() %}' in text:
            print(f"找到范围二循环，位置: {i}")
            paragraphs_to_process.append(('scope2_loop', i))

        # 找到需要替换的输出段落
        if '{{ source_name }}：{{ method }}' in text:
            print(f"找到简单输出段落，位置: {i}")
            paragraphs_to_process.append(('simple_output', i))

        # 找到范围三引用
        if '{{ quantification_methods.scope_3[cat_key] }}' in text:
            print(f"找到范围三引用，位置: {i}")
            paragraphs_to_process.append(('scope3_output', i))

        i += 1

    print(f"\\n找到 {len(paragraphs_to_process)} 个需要处理的位置")

    # 由于直接删除和插入段落比较复杂，我们使用 XML 操作
    # 让我们采取更简单的方法：直接替换文本内容

    for para in doc.paragraphs:
        text = para.text

        # 替换循环变量名
        if '{% for source_name, method in quantification_methods.scope_1.items() %}' in text:
            # 需要完全替换段落内容
            clear_paragraph_runs(para)
            add_run_to_paragraph(para, '{% for source_key, method_info in quantification_methods.scope_1.items() %}')

        if '{% for source_name, method in quantification_methods.scope_2.items() %}' in text:
            clear_paragraph_runs(para)
            add_run_to_paragraph(para, '{% for source_key, method_info in quantification_methods.scope_2.items() %}')

        # 替换简单输出为结构化输出
        if '{{ source_name }}：{{ method }}' in text:
            # 这个需要替换为多行输出
            # 注意：需要反向插入，因为 insert_paragraph_after 会按顺序插入
            clear_paragraph_runs(para)
            add_run_to_paragraph(para, '排放源：{{ method_info.name }}')
            # 反向添加后续内容
            insert_paragraph_after(para, '（3）排放因子EF：{{ method_info.ef }}')
            insert_paragraph_after(para, '（2）活动数据AD：{{ method_info.ad }}')
            insert_paragraph_after(para, '（1）量化模型：{{ method_info.model }}')

        # 替换范围三引用
        if '{{ quantification_methods.scope_3[cat_key] }}' in text:
            clear_paragraph_runs(para)
            add_run_to_paragraph(para, '排放源：{{ quantification_methods.scope_3[cat_key].name }}')
            # 反向添加
            insert_paragraph_after(para, '（3）排放因子EF：{{ quantification_methods.scope_3[cat_key].ef }}')
            insert_paragraph_after(para, '（2）活动数据AD：{{ quantification_methods.scope_3[cat_key].ad }}')
            insert_paragraph_after(para, '（1）量化模型：{{ quantification_methods.scope_3[cat_key].model }}')

    # 处理表格中的内容
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    text = para.text

                    if '{% for source_name, method in quantification_methods.scope_1.items() %}' in text:
                        clear_paragraph_runs(para)
                        add_run_to_paragraph(para, '{% for source_key, method_info in quantification_methods.scope_1.items() %}')

                    if '{% for source_name, method in quantification_methods.scope_2.items() %}' in text:
                        clear_paragraph_runs(para)
                        add_run_to_paragraph(para, '{% for source_key, method_info in quantification_methods.scope_2.items() %}')

                    if '{{ source_name }}：{{ method }}' in text:
                        clear_paragraph_runs(para)
                        add_run_to_paragraph(para, '排放源：{{ method_info.name }}')

                    if '{{ quantification_methods.scope_3[cat_key] }}' in text:
                        clear_paragraph_runs(para)
                        add_run_to_paragraph(para, '排放源：{{ quantification_methods.scope_3[cat_key].name }}')

    # 保存新文档
    doc.save('template_v2.docx')
    print("已创建 template_v2.docx")

    return True


def clear_paragraph_runs(para):
    """清除段落中的所有 run"""
    for run in para.runs:
        run.text = ''


def add_run_to_paragraph(para, text):
    """向段落添加文本 run"""
    run = para.add_run(text)


def insert_paragraph_after(para, text):
    """在指定段落后插入新段落"""
    # 获取段落的父元素
    parent = para._element.getparent()

    # 获取当前段落的索引
    index = parent.index(para._element)

    # 创建新段落
    new_para_element = OxmlElement('w:p')
    new_para = type(para)(new_para_element, para._parent)

    # 添加文本
    new_para.add_run(text)

    # 插入到父元素中
    parent.insert(index + 1, new_para_element)


if __name__ == '__main__':
    try:
        upgrade_template_v2()
        print("\\n升级完成！")
        print("新文件: template_v2.docx")
    except Exception as e:
        print(f"升级失败: {e}")
        import traceback
        traceback.print_exc()
