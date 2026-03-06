# -*- coding: utf-8 -*-
"""
重构 template.docx 的量化方法说明部分
目标：将所有 Jinja2 标签紧凑化，消除多余空行
"""
from docx import Document
from docx.oxml import OxmlElement
import shutil

def remove_paragraph(para):
    """安全删除段落"""
    parent = para._element.getparent()
    parent.remove(para._element)

def fix_template_layout():
    """重构模板布局"""
    # 备份
    shutil.copy('template.docx', 'template_backup_before_fix.docx')
    print("已备份template.docx")

    doc = Document('template.docx')

    print('=== 开始重构量化方法说明部分 ===')

    # 找到量化方法说明章节
    start_idx = None
    for i, para in enumerate(doc.paragraphs):
        if '量化方法说明' in para.text:
            start_idx = i
            break

    if not start_idx:
        print("未找到量化方法说明章节")
        return

    print(f"找到量化方法说明章节，位置: {start_idx}")

    # 新的量化方法说明部分的完整内容
    new_content = [
        # 量化方法说明标题
        '量化方法说明',

        # 范围一
        '{% set ns_scope1 = namespace(count=1) %}{% if flags.has_scope_1 %}',
        '范围一：直接温室气体排放',
        '{% for source_key, method_info in quantification_methods.scope_1.items() %}{% if loop.first %}{% else %}\n{% endif %}（{{ cn_nums[ns_scope1.count] }}）排放源：{{ method_info.name }}\n（1）量化模型：{{ method_info.model }}\n（2）活动数据AD：{{ method_info.ad }}\n（3）排放因子EF：{{ method_info.ef }}{% set ns_scope1.count = ns_scope1.count + 1 %}{% endfor %}{% endif %}',

        # 范围二（基于位置）
        '{% set ns_scope2_loc = namespace(count=1) %}{% if flags.has_scope_2_location %}',
        '范围二：外购电力（基于位置的间接温室气体排放）',
        '{% for source_key, method_info in quantification_methods.scope_2.items() %}（{{ cn_nums[ns_scope2_loc.count] }}）排放源：{{ method_info.name }}\n（1）量化模型：{{ method_info.model }}\n（2）活动数据AD：{{ method_info.ad }}\n（3）排放因子EF：{{ method_info.ef }}{% set ns_scope2_loc.count = ns_scope2_loc.count + 1 %}{% endfor %}{% endif %}',

        # 范围二（基于市场）
        '{% set ns_scope2_mkt = namespace(count=1) %}{% if flags.has_scope_2_market %}',
        '范围二：外购电力（基于市场的间接温室气体排放）',
        '{% for source_key, method_info in quantification_methods.scope_2.items() %}（{{ cn_nums[ns_scope2_mkt.count] }}）排放源：{{ method_info.name }}\n（1）量化模型：{{ method_info.model }}\n（2）活动数据AD：{{ method_info.ad }}\n（3）排放因子EF：{{ method_info.ef }}{% set ns_scope2_mkt.count = ns_scope2_mkt.count + 1 %}{% endfor %}{% endif %}',

        # 范围三
        '{% set ns_scope3 = namespace(count=1) %}{% if flags.has_scope_3 %}',
        '范围三：其他间接温室气体排放',
        '{% for i in range(1, 16) %}{% set cat_key = "category_" ~ i %}{% set flag_key = "has_scope_3_category_" ~ i %}{% if flags[flag_key] %}（{{ cn_nums[ns_scope3.count] }}）{{ scope_3_category_names["category_" ~ i] }}\n（1）量化模型：{{ quantification_methods.scope_3[cat_key].model }}\n（2）活动数据AD：{{ quantification_methods.scope_3[cat_key].ad }}\n（3）排放因子EF：{{ quantification_methods.scope_3[cat_key].ef }}{% set ns_scope3.count = ns_scope3.count + 1 %}{% endif %}{% endfor %}{% endif %}',
    ]

    # 找到量化方法说明部分的结束位置
    end_idx = start_idx + 1
    for i in range(start_idx + 1, len(doc.paragraphs)):
        text = doc.paragraphs[i].text
        # 找到下一个主要章节（如"四、"或"参考文献"等）
        if text and ('四、' in text or '参考文献' in text or '附录' in text):
            end_idx = i
            break
        # 或者找到很大段的空白后
        if i > start_idx + 100:
            end_idx = i
            break

    print(f"量化方法说明部分: {start_idx} 到 {end_idx}")

    # 删除旧的量化方法说明段落
    print(f"删除旧的量化方法说明段落（{start_idx}-{end_idx}）...")
    for i in range(end_idx - 1, start_idx - 1, -1):
        if i < len(doc.paragraphs):
            remove_paragraph(doc.paragraphs[i])

    # 在原位置插入新的内容
    print("插入新的量化方法说明内容...")
    insert_pos = start_idx

    if insert_pos < len(doc.paragraphs):
        # 在找到的位置插入
        parent = doc.paragraphs[insert_pos]._element.getparent()
        index = parent.index(doc.paragraphs[insert_pos]._element)
    else:
        parent = doc.paragraphs[-1]._element.getparent()
        index = parent.index(doc.paragraphs[-1]._element) + 1

    # 反向插入，保持顺序
    for content in reversed(new_content):
        new_para_element = OxmlElement('w:p')
        new_para = type(doc.paragraphs[0])(new_para_element, doc)

        # 处理换行符，创建多个run
        lines = content.split('\n')
        for j, line in enumerate(lines):
            if j > 0:
                # 添加换行
                new_para.add_run('\n')
            new_para.add_run(line)

        parent.insert(index, new_para_element)

    print("量化方法说明部分重构完成")

    # 检查表3和范围三的分隔
    print("\n=== 检查表3和范围三的分隔 ===")
    for i, para in enumerate(doc.paragraphs):
        text = para.text
        if '表 3' in text or 'GWP' in text:
            print(f"找到表3相关内容: 行{i} - {text[:50]}")
            # 检查后续几行是否有范围三
            for j in range(i + 1, min(i + 10, len(doc.paragraphs))):
                if '范围三' in doc.paragraphs[j].text:
                    print(f"  范围三在表3后{j-i}行处")
                    break

    # 保存
    doc.save('template_fixed.docx')
    print("\n已保存到 template_fixed.docx")

    # 替换原文件
    if os.path.exists('template.docx'):
        os.remove('template.docx')
    shutil.move('template_fixed.docx', 'template.docx')
    print("已替换 template.docx")

if __name__ == '__main__':
    import os
    try:
        fix_template_layout()
        print("\n=== 模板重构完成 ===")
    except Exception as e:
        print(f"错误: {e}")
        import traceback
        traceback.print_exc()
