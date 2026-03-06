# -*- coding: utf-8 -*-
"""
更新模板，添加自动编号系统和空白控制
"""
import sys
from docx import Document


def update_template_with_auto_numbering():
    """
    更新 template.docx，添加自动编号系统和空白控制
    """
    print("开始更新 template.docx...")

    doc = Document('template.docx')

    # 需要更新的内容映射
    replacements = [
        # 范围一：添加计数器
        (
            '{% if flags.has_scope_1 %}\n范围一：直接温室气体排放',
            '{%- if flags.has_scope_1 -%}\n范围一：直接温室气体排放\n{% set ns = namespace(count=1) %}'
        ),

        # 范围一循环：添加空白控制和编号
        (
            '{% for source_key, method_info in quantification_methods.scope_1.items() %}',
            '{%- for source_key, method_info in quantification_methods.scope_1.items() -%}'
        ),

        # 范围一排放源标题：添加编号
        (
            '排放源：{{ method_info.name }}',
            '（{{ ns.count | cn_num }}）排放源：{{ method_info.name }}'
        ),

        # 范围一结束：添加计数器递增
        (
            '（3）排放因子EF：{{ method_info.ef }}',
            '（3）排放因子EF：{{ method_info.ef }}\n{% set ns.count = ns.count + 1 %}'
        ),

        # 范围一 endif：添加空白控制
        (
            '{% endif %}\n{% if flags.has_scope_2_location %}',
            '{%- endif %}\n{%- if flags.has_scope_2_location -%}\n{% set ns = namespace(count=1) %}'
        ),

        # 范围二循环
        (
            '{% for source_key, method_info in quantification_methods.scope_2.items() %}',
            '{%- for source_key, method_info in quantification_methods.scope_2.items() -%}'
        ),

        # 范围二排放源标题
        (
            '排放源：{{ method_info.name }}',
            '（{{ ns.count | cn_num }}）排放源：{{ method_info.name }}'
        ),

        # 范围二结束：添加计数器递增
        (
            '（3）排放因子EF：{{ method_info.ef }}',
            '（3）排放因子EF：{{ method_info.ef }}\n{% set ns.count = ns.count + 1 %}'
        ),

        # 范围三：添加计数器
        (
            '{% if flags.has_scope_3 %}\n范围三：其他间接温室气体排放',
            '{%- if flags.has_scope_3 -%}\n范围三：其他间接温室气体排放\n{% set ns = namespace(count=1) %}'
        ),

        # 范围三循环
        (
            '{% for i in range(1, 16) %}',
            '{%- for i in range(1, 16) -%}'
        ),

        # 范围三内层 if
        (
            '{% if flags[flag_key] %}',
            '{%- if flags[flag_key] -%}'
        ),

        # 范围三排放源标题
        (
            '排放源：{{ quantification_methods.scope_3[cat_key].name }}',
            '（{{ ns.count | cn_num }}）排放源：{{ quantification_methods.scope_3[cat_key].name }}'
        ),

        # 范围三结束：添加计数器递增和空白控制
        (
            '（3）排放因子EF：{{ quantification_methods.scope_3[cat_key].ef }}',
            '（3）排放因子EF：{{ quantification_methods.scope_3[cat_key].ef }}\n{% set ns.count = ns.count + 1 %}'
        ),

        # 范围三内层 endif
        (
            '{% endif %}\n{% endfor %}',
            '{%- endif -%}\n{% endfor -%}'
        ),
    ]

    # 执行替换
    for para in doc.paragraphs:
        for old, new in replacements:
            if old in para.text:
                for run in para.runs:
                    run.text = run.text.replace(old, new)

    # 处理表格
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    for old, new in replacements:
                        if old in para.text:
                            for run in para.runs:
                                run.text = run.text.replace(old, new)

    # 保存新模板
    doc.save('template_auto_number.docx')
    print('已创建 template_auto_number.docx')

    return True


if __name__ == '__main__':
    try:
        update_template_with_auto_numbering()
        print("\n更新完成！")
        print("新文件: template_auto_number.docx")
        print("\n请验证并替换原 template.docx")
    except Exception as e:
        print(f"更新失败: {e}")
        import traceback
        traceback.print_exc()
