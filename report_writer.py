from docx import Document
from docx.shared import Pt, Inches
from docx.oxml.ns import qn # 关键：用于处理中文字体
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime 

class WordReportWriter: 
    def __init__(self, template_path=None, cover_image_path=None): 
        """
        初始化时，创建一个空白文档，并立即设置好样式。
        
        Args:
            template_path: 模板文件路径（暂未使用）
            cover_image_path: 封面图片路径（暂未使用，当前使用默认路径）
        """
        self.doc = Document() 
        print("创建新 Word 文档") 
        self._setup_styles()
        self._setup_page_margins()
        self.template_path = template_path
        self.cover_image_path = cover_image_path 

    def _setup_styles(self): 
        """
        "样式与内容分离"思想的核心。
        在这里定义好所有"规则"。
        """
        print("开始设置 Word 样式...") 
        
        # 1. 设置默认（正文）样式 
        style = self.doc.styles['Normal'] 
        style.font.name = '宋体' 
        # 关键一步：设置东亚字体（处理中文字符） 
        style._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体') 
        style.font.size = Pt(10.5)
        # 设置段落间距
        style.paragraph_format.space_after = Pt(0)  # 段后0点
        style.paragraph_format.line_spacing = 1.0  # 单倍行距

        # 2. 设置一级标题样式 
        style = self.doc.styles['Heading 1'] 
        style.font.name = '黑体' 
        style._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体') 
        style.font.size = Pt(16)
        style.paragraph_format.space_after = Pt(12)  # 段后12点
        
        print("样式设置完成。")
        
    def _setup_page_margins(self):
        """
        设置页面边距，匹配模板文档格式。
        """
        print("设置页面边距...")
        sections = self.doc.sections
        for section in sections:
            # 设置标准边距（1英寸 = 2.54厘米）
            section.top_margin = Inches(1.25)  # 上边距
            section.bottom_margin = Inches(1.0)  # 下边距
            section.left_margin = Inches(1.25)  # 左边距
            section.right_margin = Inches(1.25)  # 右边距
        print("页面边距设置完成。")

    def _format_number(self, value, decimals=2):
        """
        格式化数字：添加千分位分隔符，保留指定小数位数
        这是展示层的格式化逻辑，数据层保持纯净的 float 类型

        Args:
            value: 数值（float 或 int）
            decimals: 小数位数，默认 2

        Returns:
            格式化后的字符串
        """
        try:
            return f"{float(value):,.{decimals}f}"
        except (ValueError, TypeError):
            return "0.00"

    def _to_float_safe(self, value, default=0.0):
        """
        安全地将值转换为 float（防御性编程）

        Args:
            value: 任意类型的值
            default: 转换失败时的默认值

        Returns:
            float 类型的值
        """
        try:
            return float(value)
        except (ValueError, TypeError):
            return default

    def _has_emission_data(self, value):
        """
        安全地检查排放数据是否大于 0
        支持字符串和数值类型的输入

        Args:
            value: 任意类型的值

        Returns:
            bool: 如果值大于 0 返回 True，否则返回 False
        """
        float_val = self._to_float_safe(value, 0.0)
        return float_val > 0 

    def add_title_page(self, company_name, report_year):
        """
        添加封面页，完全参照模板1的格式
        """
        import os
        from datetime import datetime
        from docx.shared import Pt, Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        # 添加封面图片（如果存在）
        script_dir = os.path.dirname(os.path.abspath(__file__))
        cover_image_path = os.path.join(script_dir, "封面.png")

        print(f"尝试查找封面图片: {cover_image_path}")
        print(f"封面图片是否存在: {os.path.exists(cover_image_path)}")

        # 1. 添加几个空行作为顶部留白
        for _ in range(2):
            self.doc.add_paragraph()

        # 2. 添加封面图片（居中）
        if os.path.exists(cover_image_path):
            paragraph = self.doc.add_paragraph()
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = paragraph.add_run()
            # 添加图片，设置适当的宽度（5.5英寸）
            run.add_picture(cover_image_path, width=Inches(5.5))
            print("已添加封面图片")

            # 在图片后添加一些空行
            for _ in range(2):
                self.doc.add_paragraph()
        else:
            # 如果没有图片，添加更多空行作为留白
            for _ in range(4):
                self.doc.add_paragraph()

        # 3. 公司名称 - 26pt，加粗，居中
        company_paragraph = self.doc.add_paragraph()
        company_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        company_run = company_paragraph.add_run(company_name)
        company_run.font.size = Pt(26)  # 模板中的26pt
        company_run.font.bold = True
        company_run.font.name = '黑体'
        company_run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')

        # 4. 报告标题 - 26pt，加粗，居中
        title_paragraph = self.doc.add_paragraph()
        title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title_paragraph.add_run("温室气体排放盘查报告")
        title_run.font.size = Pt(26)  # 模板中的26pt
        title_run.font.bold = True
        title_run.font.name = '黑体'
        title_run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')

        # 添加空行
        for _ in range(3):
            self.doc.add_paragraph()

        # 5. 统计期间 - 14pt，左对齐
        period_paragraph = self.doc.add_paragraph()
        period_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        period_paragraph.paragraph_format.space_after = Pt(0)  # 段后不留空
        period_run = period_paragraph.add_run(f"统计期间：{report_year}年1月1日~ {report_year}年12月31日")
        period_run.font.size = Pt(14)  # 模板中的14pt
        period_run.font.name = '宋体'
        period_run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

        # 6. 版本号 - 14pt，左对齐
        version_paragraph = self.doc.add_paragraph()
        version_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        version_paragraph.paragraph_format.space_after = Pt(0)  # 段后不留空
        version_run = version_paragraph.add_run("版本号：A/1")
        version_run.font.size = Pt(14)
        version_run.font.name = '宋体'
        version_run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

        # 7. 文件编号 - 14pt，左对齐
        file_number_paragraph = self.doc.add_paragraph()
        file_number_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        file_number_paragraph.paragraph_format.space_after = Pt(0)  # 段后不留空
        # 生成基于当前日期的文件编号
        current_date = datetime.now()
        file_number = f"DY-GHG-{current_date.strftime('%Y')}-01"
        file_number_run = file_number_paragraph.add_run(f"文件编号：{file_number}")
        file_number_run.font.size = Pt(14)
        file_number_run.font.name = '宋体'
        file_number_run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

        # 8. 编制日期 - 14pt，左对齐
        create_date_paragraph = self.doc.add_paragraph()
        create_date_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        create_date_paragraph.paragraph_format.space_after = Pt(0)  # 段后不留空
        create_date_run = create_date_paragraph.add_run(f"编制日期：{current_date.strftime('%Y年%m月%d日')}")
        create_date_run.font.size = Pt(14)
        create_date_run.font.name = '宋体'
        create_date_run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

        # 9. 修订日期 - 14pt，左对齐
        revise_date_paragraph = self.doc.add_paragraph()
        revise_date_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        revise_date_paragraph.paragraph_format.space_after = Pt(0)  # 段后不留空
        revise_date_run = revise_date_paragraph.add_run(f"修订日期：{current_date.strftime('%Y年%m月%d日')}")
        revise_date_run.font.size = Pt(14)
        revise_date_run.font.name = '宋体'
        revise_date_run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

        # 添加大量空行将签名区域推到页面底部
        for _ in range(8):
            self.doc.add_paragraph()

        # 10. 签名区域标题 - 14pt，居中
        signature_header_paragraph = self.doc.add_paragraph()
        signature_header_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        signature_header_run = signature_header_paragraph.add_run("审核人：                          编制人：                          批准人：")
        signature_header_run.font.size = Pt(14)
        signature_header_run.font.name = '宋体'
        signature_header_run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

        # 11. 底部信息 - 统一设置对齐和段落格式

        # 审核信息
        review_paragraph = self.doc.add_paragraph()
        review_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        review_paragraph.paragraph_format.space_after = Pt(0)  # 段后不留空
        review_run = review_paragraph.add_run("审核")
        review_run.font.size = Pt(14)
        review_run.font.bold = True
        review_run.font.name = '宋体'
        review_run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        review_paragraph.paragraph_format.line_spacing = 1.5

        # 组织
        org_paragraph = self.doc.add_paragraph()
        org_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        org_paragraph.paragraph_format.space_after = Pt(0)  # 段后不留空
        org_run = org_paragraph.add_run("审核组织")
        org_run.font.size = Pt(12)  # 默认字体大小
        org_run.font.name = '宋体'
        org_run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        org_paragraph.paragraph_format.line_spacing = 1.5

        # 公司名称
        company_info_paragraph = self.doc.add_paragraph()
        company_info_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        company_info_paragraph.paragraph_format.space_after = Pt(0)  # 段后不留空
        company_info_run = company_info_paragraph.add_run(f"公司名称：{company_name}")
        company_info_run.font.size = Pt(12)
        company_info_run.font.bold = True  # 标签加粗
        company_info_run.font.name = '宋体'
        company_info_run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        company_info_paragraph.paragraph_format.line_spacing = 1.5

        # 统一社会信用代码（使用示例代码）
        credit_code_paragraph = self.doc.add_paragraph()
        credit_code_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        credit_code_paragraph.paragraph_format.space_after = Pt(0)  # 段后不留空
        credit_code_run = credit_code_paragraph.add_run("统一社会信用代码：91420000798750168P")
        credit_code_run.font.size = Pt(12)
        credit_code_run.font.bold = True  # 标签加粗
        credit_code_run.font.name = '宋体'
        credit_code_run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        credit_code_paragraph.paragraph_format.line_spacing = 1.5

        # 编制人
        preparer_paragraph = self.doc.add_paragraph()
        preparer_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        preparer_paragraph.paragraph_format.space_after = Pt(0)  # 段后不留空
        preparer_run = preparer_paragraph.add_run("编制人：系统自动生成")
        preparer_run.font.size = Pt(12)
        preparer_run.font.bold = True  # 标签加粗
        preparer_run.font.name = '宋体'
        preparer_run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        preparer_paragraph.paragraph_format.line_spacing = 1.5

        print(f"已添加模板格式的封面: {company_name}")

        # 添加分页符
        self.doc.add_page_break() 

    def add_executive_summary(self, summary):
        """
        添加执行摘要部分到报告中
        
        Args:
            summary: 执行摘要文本内容
        """
        print("开始添加执行摘要...")
        self.doc.add_heading("执行摘要", level=1)
        
        # 添加摘要段落
        paragraph = self.doc.add_paragraph()
        paragraph.add_run(summary)
        
        # 添加分页符
        self.doc.add_page_break()
        print("执行摘要添加完成。")
        
    def add_emission_table(self, data):
        """
        接收数据字典，动态生成排放汇总表格（积木模式）。
        根据 flags 系统和实际数据动态生成表格行，不写死行数。

        Args:
            data: 包含排放数据和 flags 的字典
        """
        print("开始生成排放汇总表...")
        self.doc.add_heading("1. 排放数据汇总", level=1)

        # 创建表格（1行表头 + 动态数据行, 3列：项目、排放量、备注）
        table = self.doc.add_table(rows=1, cols=3)
        table.style = 'Table Grid'

        # 填充表头
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = '排放项目'
        hdr_cells[1].text = '排放量(tCO2e)'
        hdr_cells[2].text = '备注'

        # 设置表头样式
        for i in range(3):
            hdr_cells[i].paragraphs[0].runs[0].bold = True

        # 获取 flags
        flags = data.get('flags', {})

        # ========== 动态构建表格数据（积木模式）==========
        # 每个元组格式：(项目名称, 排放量值, 备注, 是否显示的条件)
        table_items = []

        # 1. 总排放量（总是显示）
        table_items.append((
            "总排放量(基于位置)",
            data.get('total_emission_location', 0.0),
            "范围一+范围二(基于位置)+范围三",
            True
        ))
        table_items.append((
            "总排放量(基于市场)",
            data.get('total_emission_market', 0.0),
            "范围一+范围二(基于市场)+范围三",
            True
        ))

        # 2. 范围一（根据 flags 决定是否显示）
        table_items.append((
            "范围一：直接温室气体排放",
            data.get('scope_1_emissions', 0.0),
            "固定燃烧、移动燃烧、散逸排放、工艺过程",
            flags.get('has_scope_1', True)
        ))

        # 3. 范围二（根据 flags 决定是否显示）
        table_items.append((
            "范围二：能源间接温室气体排放(基于位置)",
            data.get('scope_2_location_based_emissions', 0.0),
            "外购电力、热力、蒸汽等",
            flags.get('has_scope_2_location', True)
        ))
        table_items.append((
            "范围二：能源间接温室气体排放(基于市场)",
            data.get('scope_2_market_based_emissions', 0.0),
            "外购电力、热力、蒸汽等",
            flags.get('has_scope_2_market', True)
        ))

        # 4. 范围三（根据 flags 决定是否显示）
        table_items.append((
            "范围三：其他间接温室气体排放",
            data.get('scope_3_emissions', 0.0),
            "上游运输、商务差旅、员工通勤等",
            flags.get('has_scope_3', True)
        ))

        # ========== 动态添加行（循环而非硬编码）==========
        for name, value, note, should_show in table_items:
            if should_show:
                row_cells = table.add_row().cells
                row_cells[0].text = name
                row_cells[1].text = self._format_number(value)
                row_cells[2].text = note

        print(f"表格生成完毕，共 {len(table.rows)} 行（含表头）。")

    def add_emission_source_table(self, title, items):
        """
        动态添加排放源明细表格（积木模式）。
        根据 items 列表动态生成表格行，支持任意数量的排放源。

        Args:
            title: 表格标题
            items: 排放源列表，每个元素是包含 name, emission, note 的字典
        """
        if not items:
            print(f"没有排放源数据，跳过 {title}")
            return

        print(f"开始生成 {title} 表格，共 {len(items)} 个排放源...")
        self.doc.add_heading(title, level=2)

        # 创建表格（1行表头 + 动态数据行, 3列）
        table = self.doc.add_table(rows=1, cols=3)
        table.style = 'Table Grid'

        # 填充表头
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = '类别/排放源'
        hdr_cells[1].text = '排放量(tCO2e)'
        hdr_cells[2].text = '备注/设施'

        # 设置表头样式
        for i in range(3):
            hdr_cells[i].paragraphs[0].runs[0].bold = True

        # ========== 动态添加排放源行（循环）==========
        for item in items:
            row_cells = table.add_row().cells
            row_cells[0].text = item.get('name', '')
            # 格式化数值
            emission_value = item.get('emission', 0.0)
            if isinstance(emission_value, (int, float)):
                row_cells[1].text = self._format_number(emission_value)
            else:
                row_cells[1].text = str(emission_value)
            row_cells[2].text = item.get('note', '')

        print(f"{title} 表格生成完毕，共 {len(table.rows)} 行（含表头）。")

    def _add_no_data_explanation(self, section_title, main_explanation, background_info):
        """
        生成专业的不涉及排放说明文字

        Args:
            section_title: 章节标题
            main_explanation: 主要说明文字（如"经核查，本年度未涉及..."）
            background_info: 背景信息/定义解释
        """
        # 添加主要说明段落（加粗）
        paragraph = self.doc.add_paragraph()
        run = paragraph.add_run(main_explanation)
        run.font.bold = True

        # 添加背景信息段落（普通文本）
        paragraph = self.doc.add_paragraph()
        paragraph.add_run(background_info)

        print(f"已添加 {section_title} 的无数据说明文字。")

    def _add_scope_1_chapter(self, data):
        """
        添加范围一章节（固定燃烧、移动燃烧、散逸排放、工艺过程）

        Args:
            data: 包含排放数据的字典
        """
        flags = data.get('flags', {})
        has_scope_1 = flags.get('has_scope_1', False)

        print("开始添加范围一章节...")
        self.doc.add_heading("2. 范围一：直接温室气体排放", level=1)

        if has_scope_1:
            # 有数据：添加排放源表格
            scope1_items = data.get('scope1_items', [])
            if scope1_items:
                self.add_emission_source_table("范围一排放源明细", scope1_items)
            else:
                # 有数值但无明细，添加描述性文本
                paragraph = self.doc.add_paragraph()
                paragraph.add_run("范围一排放包括固定燃料燃烧、移动燃料燃烧、散逸排放和工艺过程排放等直接温室气体排放。")
        else:
            # 无数据：生成专业说明文字
            self._add_no_data_explanation(
                "范围一：直接温室气体排放",
                "经核查，本年度未涉及范围一相关排放活动（或数据不具备重要性），故不进行量化。",
                "范围一排放指企业拥有或控制的排放源产生的直接温室气体排放，包括固定燃料燃烧（如锅炉、熔炉）、移动燃料燃烧（如车辆、船舶）、散逸排放（如制冷剂泄漏）和工艺过程排放等。"
            )

        print("范围一章节添加完成。")

    def _add_scope_2_chapter(self, data):
        """
        添加范围二章节（能源间接温室气体排放）

        Args:
            data: 包含排放数据的字典
        """
        flags = data.get('flags', {})
        has_scope_2 = flags.get('has_scope_2_location', False) or flags.get('has_scope_2_market', False)

        print("开始添加范围二章节...")
        self.doc.add_heading("3. 范围二：能源间接温室气体排放", level=1)

        if has_scope_2:
            # 有数据：添加描述和数值
            paragraph = self.doc.add_paragraph()
            paragraph.add_run("范围二排放指企业外购的电力、热力、蒸汽等能源产生的间接温室气体排放。")

            scope2_location = data.get('scope_2_location_based_emissions', 0.0)
            scope2_market = data.get('scope_2_market_based_emissions', 0.0)

            # 使用安全的数值检查方法
            if self._has_emission_data(scope2_location):
                paragraph = self.doc.add_paragraph()
                paragraph.add_run(f"基于位置法的范围二排放量为：{self._format_number(scope2_location)} tCO2e")

            if self._has_emission_data(scope2_market):
                paragraph = self.doc.add_paragraph()
                paragraph.add_run(f"基于市场法的范围二排放量为：{self._format_number(scope2_market)} tCO2e")
        else:
            # 无数据：生成专业说明文字
            self._add_no_data_explanation(
                "范围二：能源间接温室气体排放",
                "经核查，本年度未涉及范围二相关排放活动（或数据不具备重要性），故不进行量化。",
                "范围二排放指企业外购的电力、热力、蒸汽等能源在生产、运输过程中产生的间接温室气体排放。计算方法可分为基于位置法（使用电网平均排放因子）和基于市场法（使用特定电力合同排放因子）。"
            )

        print("范围二章节添加完成。")

    def _add_scope_3_chapter(self, data):
        """
        添加范围三章节（其他间接温室气体排放）

        Args:
            data: 包含排放数据的字典
        """
        flags = data.get('flags', {})
        has_scope_3 = flags.get('has_scope_3', False)

        print("开始添加范围三章节...")
        self.doc.add_heading("4. 范围三：其他间接温室气体排放", level=1)

        if has_scope_3:
            # 有数据：添加排放源表格
            scope2_3_items = data.get('scope2_3_items', [])
            if scope2_3_items:
                self.add_emission_source_table("范围三排放源明细", scope2_3_items)
            else:
                # 有数值但无明细，添加描述性文本
                paragraph = self.doc.add_paragraph()
                paragraph.add_run("范围三排放包括上游运输、商务差旅、员工通勤、产品使用等价值链上下游的其他间接温室气体排放。")
        else:
            # 无数据：生成专业说明文字（根据用户要求）
            self._add_no_data_explanation(
                "范围三：其他间接温室气体排放",
                "经核查，本年度未涉及范围三相关排放类别（或数据不具备重要性），故不进行量化。",
                "范围三排放指价值链上游和下游的其他间接温室气体排放，包括但不限于：外购商品和服务、资本货物、燃料和能源相关活动、上下游运输配送、运营中产生的废弃物、商务差旅、员工通勤、已售产品的使用和报废处理等。根据《温室气体核算体系》，范围三包含15个类别，企业可根据实际情况选择性量化。"
            )

        print("范围三章节添加完成。")

    def save(self, output_path): 
        """
        保存最终生成的 Word 文档，添加错误处理以处理文件锁定等情况。
        """
        try:
            # 尝试直接保存
            self.doc.save(output_path)
            print(f"文档已成功保存到: {output_path}")
            return True
        except PermissionError:
            # 如果文件被锁定，尝试使用时间戳重命名保存
            import os
            from datetime import datetime
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            base_name, ext = os.path.splitext(output_path)
            new_output_path = f"{base_name}_{timestamp}{ext}"
            self.doc.save(new_output_path)
            print(f"原文件被锁定，已保存为: {new_output_path}")
            return True
        except Exception as e:
            print(f"保存文件失败: {e}")
            return False

# (你可以在文件末尾添加测试代码)
    def write_report(self, data, output_path):
        """
        主方法：整合所有功能，生成完整报告（组件化 + 条件判断）

        Args:
            data: 包含报告所有数据的字典（应包含 flags）
            output_path: 输出文件路径
        """
        print(f"开始生成完整报告: {output_path}")

        # ========== 断点调试：检查数据类型 ==========
        print("\n========== 数据类型调试信息 ==========")
        if 'greenhouse_gas_data' in data:
            ghg_data = data['greenhouse_gas_data']
        else:
            ghg_data = data

        # 打印关键字段的数据类型和值
        debug_keys = [
            'scope_1_emissions',
            'scope_2_location_based_emissions',
            'scope_2_market_based_emissions',
            'scope_3_emissions',
            'total_emission_location',
            'total_emission_market'
        ]
        for key in debug_keys:
            value = ghg_data.get(key, 'NOT_FOUND')
            value_type = type(value).__name__
            print(f"  {key}: {value_type} = {value}")

        # 检查 flags
        flags = ghg_data.get('flags', {})
        print(f"  flags: {flags}")
        print("======================================\n")

        # ========== 步骤1：提取主要数据（兼容新旧数据结构）==========
        if 'greenhouse_gas_data' in data:
            ghg_data = data['greenhouse_gas_data']
            emission_reductions = data.get('emission_reductions', [])
            executive_summary = data.get('executive_summary', ghg_data.get('executive_summary'))
        else:
            ghg_data = data
            emission_reductions = data.get('emission_reductions', [])
            executive_summary = data.get('executive_summary')

        # 确保 ghg_data 有 flags（向后兼容）
        if 'flags' not in ghg_data:
            # 如果没有 flags，创建默认的 flags（假设所有数据都存在）
            ghg_data['flags'] = {
                'has_scope_1': ghg_data.get('scope_1_emissions', 0) > 0,
                'has_scope_2_location': ghg_data.get('scope_2_location_based_emissions', 0) > 0,
                'has_scope_2_market': ghg_data.get('scope_2_market_based_emissions', 0) > 0,
                'has_scope_3': ghg_data.get('scope_3_emissions', 0) > 0,
            }

        # ========== 步骤2：添加封面页（始终生成）==========
        company_name = ghg_data.get('company_name', '未知公司')
        report_year = ghg_data.get('report_year', datetime.now().year)
        self.add_title_page(company_name, report_year)

        # ========== 步骤3：添加执行摘要（如果有）==========
        if executive_summary:
            self.add_executive_summary(executive_summary)
        else:
            print("警告：未找到执行摘要数据")
            self.doc.add_page_break()

        # ========== 步骤4：添加排放数据汇总表（动态生成）==========
        self.add_emission_table(ghg_data)

        # ========== 步骤5：添加各范围章节（智能业务逻辑层）==========
        # 注意：现在这些方法内部会处理"有数据"和"无数据"两种情况
        # 无数据时会自动生成专业说明文字，而非跳过章节
        self._add_scope_1_chapter(ghg_data)
        self._add_scope_2_chapter(ghg_data)
        self._add_scope_3_chapter(ghg_data)

        # ========== 步骤6：保存文档 ==========
        return self.save(output_path)

# (你可以在文件末尾添加测试代码) 
if __name__ == "__main__": 
    from datetime import datetime
    writer = WordReportWriter() 
    test_data = { 
        'company_name': '我的测试公司',
        'report_year': 2024,
        'total_emission': 1000, 
        'scope_1': 600, 
        'scope_2': 400,
        'executive_summary': '这是一个测试用的执行摘要，展示了AI生成的碳盘查报告执行摘要的基本格式和内容。'
    } 
    writer.write_report(test_data, "test_report.docx") 
    print("--- 测试 report_writer.py 完成 ---")