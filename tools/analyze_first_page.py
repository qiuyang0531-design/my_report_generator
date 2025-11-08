from docx import Document
import os

def analyze_first_page():
    """
    详细分析模板1.docx第一页的格式
    """
    template_path = "模板1.docx"

    if not os.path.exists(template_path):
        print(f"模板文件 {template_path} 不存在")
        return

    try:
        doc = Document(template_path)
        print("===== 模板1第一页详细分析 =====")

        # 分析第一页的内容（前20个段落，大概在前几页）
        first_page_paragraphs = []
        content_count = 0

        for i, paragraph in enumerate(doc.paragraphs):
            if paragraph.text.strip():  # 只关注有内容的段落
                content_count += 1
                first_page_paragraphs.append((i, paragraph))
                # 收集前15个有内容的段落，这应该能覆盖第一页
                if content_count >= 15:
                    break

        print(f"第一页相关内容段落数: {len(first_page_paragraphs)}")

        for i, (para_index, paragraph) in enumerate(first_page_paragraphs):
            print(f"\n--- 内容段落 {i+1} (原文档第{para_index+1}段) ---")
            print(f"文本内容: '{paragraph.text}'")

            if paragraph.runs:
                for j, run in enumerate(paragraph.runs):
                    if run.text.strip():
                        print(f"  运行 {j+1}: '{run.text}'")
                        print(f"    字体: {run.font.name if run.font.name else '默认'}")
                        print(f"    字号: {run.font.size.pt if run.font.size else '默认'} pt")
                        print(f"    加粗: {run.font.bold}")
                        print(f"    斜体: {run.font.italic}")
                        print(f"    颜色: {run.font.color.rgb if run.font.color.rgb else '自动'}")

                print(f"  段落对齐: {paragraph.alignment}")

                # 分析段前段后间距
                if paragraph.paragraph_format.space_before:
                    print(f"  段前间距: {paragraph.paragraph_format.space_before.pt} pt")
                if paragraph.paragraph_format.space_after:
                    print(f"  段后间距: {paragraph.paragraph_format.space_after.pt}")
                if paragraph.paragraph_format.line_spacing:
                    print(f"  行间距: {paragraph.paragraph_format.line_spacing}")

        # 专门检查图片位置
        print(f"\n===== 图片位置分析 =====")
        for i, paragraph in enumerate(doc.paragraphs):
            if paragraph.text.strip() == "":
                # 检查空段落是否包含图片
                for run in paragraph.runs:
                    if hasattr(run.element, './/pic:pic') or run.element.xpath('.//pic:pic'):
                        print(f"发现图片在段落 {i+1} (空段落)")

    except Exception as e:
        print(f"分析模板时出错: {e}")

if __name__ == "__main__":
    analyze_first_page()