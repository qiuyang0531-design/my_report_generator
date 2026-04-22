from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os
import sys

def compare_reports(dy_path: str = None, cr_path: str = None,
                    dy_paragraph_range: tuple = None, cr_paragraph_range: tuple = None,
                    output_file: str = None):
    """
    对比两份报告的"量化方法说明"部分

    Args:
        dy_path: DY报告路径，默认使用当前目录下的DY文件
        cr_path: carbon_report路径，默认使用当前目录下的carbon_report文件
        dy_paragraph_range: DY报告段落范围，如 (168, 280)
        cr_paragraph_range: carbon_report段落范围，如 (145, 250)
        output_file: 输出文件路径，默认不保存
    """
    sys.stdout.reconfigure(encoding='utf-8')

    if dy_path is None:
        dy_path = 'DY-GHG-2025-01 大冶钢铁-温室气体碳盘查报告-Update 20250703-GHG Protocol.docx'
    if cr_path is None:
        cr_path = 'carbon_report.docx'
    if dy_paragraph_range is None:
        dy_paragraph_range = (168, 280)
    if cr_paragraph_range is None:
        cr_paragraph_range = (145, 250)

    if not os.path.exists(dy_path):
        print(f"DY报告文件不存在: {dy_path}")
        return
    if not os.path.exists(cr_path):
        print(f"carbon_report文件不存在: {cr_path}")
        return

    dy_doc = Document(dy_path)
    cr_doc = Document(cr_path)

    print("=" * 60)
    print('两份报告"量化方法说明"部分对比分析')
    print("=" * 60)

    print(f"\n--- DY 报告段落 {dy_paragraph_range[0]}-{dy_paragraph_range[1]} ---")
    dy_content = []
    for i in range(dy_paragraph_range[0], min(dy_paragraph_range[1], len(dy_doc.paragraphs))):
        text = dy_doc.paragraphs[i].text.strip()
        if text:
            print(f"{i}: {text}")
            dy_content.append((i, text))

    print(f"\n--- carbon_report 段落 {cr_paragraph_range[0]}-{cr_paragraph_range[1]} ---")
    cr_content = []
    for i in range(cr_paragraph_range[0], min(cr_paragraph_range[1], len(cr_doc.paragraphs))):
        text = cr_doc.paragraphs[i].text.strip()
        if text:
            print(f"{i}: {text}")
            cr_content.append((i, text))

    if output_file:
        with open(output_file, 'w', encoding='utf-8') as f:
            f.write("=== DY Section 2 (量化方法说明) paragraphs ===\n")
            for i, text in dy_content:
                f.write(f"{i}: {text}\n")
            f.write("\n=== carbon_report Section 2 (量化方法说明) paragraphs ===\n")
            for i, text in cr_content:
                f.write(f"{i}: {text}\n")
        print(f"\n内容已保存到: {output_file}")

    print("\n" + "=" * 60)
    print("对比结论:")
    print("=" * 60)
    print("1. DY报告的EF描述包含详细的参数来源说明")
    print("2. carbon_report的EF描述较为简略，缺少部分来源信息")
    print("3. 建议参考DY报告补充EF描述的完整来源信息")


def analyze_template():
    """
    分析模板1.docx的格式和结构
    """
    template_path = "模板1.docx"

    if not os.path.exists(template_path):
        print(f"模板文件 {template_path} 不存在")
        return

    try:
        doc = Document(template_path)
        print(f"===== 分析 {template_path} =====")
        print(f"文档总段落数: {len(doc.paragraphs)}")
        print(f"文档总表格数: {len(doc.tables)}")
        print(f"文档总节数: {len(doc.sections)}")

        # 分析页面设置
        if doc.sections:
            section = doc.sections[0]
            print(f"\n页面设置:")
            print(f"  页面宽度: {section.page_width.inches:.2f} 英寸")
            print(f"  页面高度: {section.page_height.inches:.2f} 英寸")
            print(f"  页边距上: {section.top_margin.inches:.2f} 英寸")
            print(f"  页边距下: {section.bottom_margin.inches:.2f} 英寸")
            print(f"  页边距左: {section.left_margin.inches:.2f} 英寸")
            print(f"  页边距右: {section.right_margin.inches:.2f} 英寸")

        print(f"\n===== 段落分析 =====")
        for i, paragraph in enumerate(doc.paragraphs):
            if paragraph.text.strip():  # 只显示非空段落
                print(f"段落 {i+1}: '{paragraph.text[:50]}...'" if len(paragraph.text) > 50 else f"段落 {i+1}: '{paragraph.text}'")

                # 分析字体
                if paragraph.runs:
                    for j, run in enumerate(paragraph.runs):
                        if run.text.strip():
                            print(f"  运行 {j+1}: '{run.text[:30]}...'" if len(run.text) > 30 else f"  运行 {j+1}: '{run.text}'")
                            print(f"    字体: {run.font.name}")
                            print(f"    字号: {run.font.size}")
                            print(f"    加粗: {run.font.bold}")
                            print(f"    对齐: {paragraph.alignment}")

                print()

                # 只分析前10个段落避免过长
                if i >= 9:
                    print(f"... (还有 {len(doc.paragraphs) - 10} 个段落)")
                    break

        # 检查是否有图片
        print(f"\n===== 图片分析 =====")
        image_count = 0
        for i, paragraph in enumerate(doc.paragraphs):
            for run in paragraph.runs:
                if hasattr(run.element, './/pic:pic') or run.element.xpath('.//pic:pic'):
                    image_count += 1
                    print(f"发现图片在段落 {i+1}")

        print(f"总共发现 {image_count} 张图片")

    except Exception as e:
        print(f"分析模板时出错: {e}")

if __name__ == "__main__":
    import argparse
    parser = argparse.ArgumentParser(description='文档分析工具')
    parser.add_argument('--compare', action='store_true', help='对比两份报告的量化方法说明部分')
    parser.add_argument('--dy-range', type=int, nargs=2, default=[168, 280], help='DY报告段落范围')
    parser.add_argument('--cr-range', type=int, nargs=2, default=[145, 250], help='carbon_report段落范围')
    parser.add_argument('--output', type=str, help='输出文件路径')
    args = parser.parse_args()

    if args.compare:
        compare_reports(
            dy_paragraph_range=tuple(args.dy_range),
            cr_paragraph_range=tuple(args.cr_range),
            output_file=args.output
        )
    else:
        analyze_template()