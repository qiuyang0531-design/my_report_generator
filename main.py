# main.py
# 使用重构后的协议驱动型数据读取器（方式1: 从新包导入）
from data_reader import ExcelDataReaderRefactored as ExcelDataReader
from docxtpl import DocxTemplate
from jinja2 import Environment
import os
import re
from docx.oxml import OxmlElement


def to_chinese_num(n):
    """
    将数字转换为中文大写数字（用于报告编号）

    Args:
        n: 数字 (1-15)

    Returns:
        中文大写数字字符串
    """
    chinese_map = {
        1: '一', 2: '二', 3: '三', 4: '四', 5: '五',
        6: '六', 7: '七', 8: '八', 9: '九', 10: '十',
        11: '十一', 12: '十二', 13: '十三', 14: '十四', 15: '十五'
    }
    return chinese_map.get(n, str(n))


def format_number(value, decimals=2, with_comma=True):
    """
    格式化数字：添加千分位分隔符，保留指定小数位数（展示层格式化）

    Args:
        value: 数值（float 或 int）
        decimals: 小数位数，默认 2
        with_comma: 是否添加千分位分隔符，默认 True

    Returns:
        格式化后的字符串
    """
    try:
        float_value = float(value)
        if with_comma:
            return f"{float_value:,.{decimals}f}"
        else:
            return f"{float_value:.{decimals}f}"
    except (ValueError, TypeError):
        return "0.00"


def prepare_context_with_formatting(context):
    """
    为 context 添加格式化后的数值版本（展示层处理）
    保持原始数据不变，添加 _formatted 后缀的格式化版本
    对于排放量为0的类别，添加说明文字

    新增：全局字符串清洗步骤，去除所有字符串值的冗余空格
    新增：初始化所有协议变量以确保 Jinja2 兼容性

    Args:
        context: 原始数据字典

    Returns:
        包含格式化字段的新字典
    """
    formatted_context = context.copy()

    # ========== 新增：初始化所有协议变量（Jinja2兼容性）==========
    # 确保所有协议变量在模板中可用，即使Excel中没有对应表格
    protocol_output_vars = {
        'pro_ef_items': [],           # 排放因子表（新模板变量名）
        'emission_factor_items': [],  # 旧变量名（向后兼容）
        'scope1_stationary_combustion_emissions_items': [],  # 固定燃烧
        'scope1_mobile_combustion_emissions_items': [],      # 移动燃烧
        'scope1_fugitive_emissions_items': [],               # 逸散排放
        'scope1_process_emissions_items': [],                # 制程排放
        'gwp_items': [],
        'ghg_inventory_items': [],
        'activity_summary_items': [],
        'uncertainty_items': [],
        'reduction_action_items': [],
    }

    for var_name, default_value in protocol_output_vars.items():
        if var_name not in formatted_context:
            formatted_context[var_name] = default_value
            # print(f"[初始化协议变量] {var_name} = {default_value}")
    # ========== 协议变量初始化结束 ==========

    # ========== 新增：全局字符串清洗步骤 ==========
    # 遍历所有值，对字符串类型执行 strip() 去除冗余空格
    def clean_strings_in_dict(d):
        """递归清洗字典中的所有字符串值"""
        if not isinstance(d, dict):
            return d
        
        cleaned = {}
        for key, value in d.items():
            if isinstance(value, str):
                # 去除首尾空格，替换多个连续空格为单个空格
                cleaned_value = re.sub(r'\s+', ' ', str(value).strip())
                cleaned[key] = cleaned_value
            elif isinstance(value, dict):
                cleaned[key] = clean_strings_in_dict(value)
            elif isinstance(value, list):
                cleaned[key] = [clean_strings_in_dict(item) if isinstance(item, dict) else 
                               (re.sub(r'\s+', ' ', str(item).strip()) if isinstance(item, str) else item)
                               for item in value]
            else:
                cleaned[key] = value
        return cleaned
    
    formatted_context = clean_strings_in_dict(formatted_context)
    # ========== 全局清洗步骤结束 ==========

    # 计算范围三类别排放总量（用于表格汇总）
    scope3_total_sum = 0
    for i in range(1, 16):
        cat_emission = context.get(f'scope_3_category_{i}_emissions', 0)
        if cat_emission and cat_emission > 0:
            scope3_total_sum += cat_emission

    formatted_context['scope3_categories_total_sum'] = scope3_total_sum
    formatted_context['scope3_categories_total_sum_formatted'] = format_number(scope3_total_sum)

    # 计算每个类别的详细项排放量总和（用于各类别表格汇总）
    # 每个类别计算所有列的总和
    emission_columns = [
        'total_green_house_gas_emissions',
        'CO2_emissions',
        'CH4_emissions',
        'N2O_emissions',
        'HFCs_emissions',
        'PFCs_emissions',
        'SFs_emissions',
        'NF3_emissions'
    ]

    # 列名映射：数据源使用SF6，模板使用SFs
    column_key_mapping = {
        'SFs_emissions': 'SF6_emissions'  # 从数据源读取SF6，但输出为SFs
    }

    for cat_num in range(1, 16):
        category_var = f'scope3_category{cat_num}'
        if category_var in context and context[category_var]:
            # 初始化各列的总和
            column_sums = {col: 0.0 for col in emission_columns}

            # 遍历该类别的所有详细项，累加各列值
            for item in context[category_var]:
                for col in emission_columns:
                    # 根据映射获取实际的数据键名
                    data_key = column_key_mapping.get(col, col)
                    emission_str = item.get(data_key, '0')
                    # 去除逗号和空格，转换为浮点数
                    emission_str = emission_str.replace(',', '').replace(' ', '')
                    try:
                        emission_value = float(emission_str)
                        column_sums[col] += emission_value
                    except (ValueError, TypeError):
                        pass

            # 为每个列创建格式化后的总和变量
            for col in emission_columns:
                sum_value = column_sums[col]
                formatted_context[f'scope3_category{cat_num}_{col}_sum'] = sum_value
                formatted_context[f'scope3_category{cat_num}_{col}_sum_formatted'] = format_number(sum_value)

    # 需要格式化的数值字段列表
    number_fields = [
        'scope_1_emissions',
        'scope_2_location_based_emissions',
        'scope_2_market_based_emissions',
        'scope_3_emissions',
        'total_emission_location',
        'total_emission_market',
    ]

    # 范围三类别字段
    for i in range(1, 16):
        number_fields.append(f'scope_3_category_{i}_emissions')

    # 别名也需要格式化
    alias_fields = [
        'scope_1',
        'scope_2_location',
        'scope_2_market',
        'scope_3',
    ]

    all_number_fields = number_fields + alias_fields

    # 添加字段别名映射（data_reader使用短字段名，模板期望长字段名）
    field_aliases = {
        'scope_1': 'scope_1_emissions',
        'scope_2_location': 'scope_2_location_based_emissions',
        'scope_2_market': 'scope_2_market_based_emissions',
        'scope_3': 'scope_3_emissions',
    }
    
    # 为每个别名创建对应的完整字段名
    for short_name, long_name in field_aliases.items():
        if short_name in context:
            formatted_context[long_name] = context[short_name]
            formatted_context[f'{long_name}_formatted'] = format_number(context[short_name])

    # 为每个数值字段创建格式化版本（跳过已经通过别名映射创建的字段）
    for field in all_number_fields:
        formatted_key = f'{field}_formatted'
        # 如果这个字段已经被格式化过了（通过别名映射），跳过
        if formatted_key in formatted_context:
            continue
        if field in context and context[field] is not None:
            formatted_context[formatted_key] = format_number(context[field])
        else:
            formatted_context[formatted_key] = format_number(0)

    # 为范围三类别添加 display 字段（处理0值显示说明文字）
    # 有数据的类别：显示格式化数字
    for i in range(1, 16):
        field = f'scope_3_category_{i}_emissions'
        value = context.get(field, 0)
        display_value = format_number(value) if value and value > 0 else ""
        formatted_context[f'{field}_display'] = display_value
        # 同时添加不带 scope3_ 前缀的别名（兼容模板）
        formatted_context[f'category_{i}_emissions_display'] = display_value

    # 收集所有无数据的范围三类别，生成汇总说明
    categories_not_in_scope = []
    for i in range(1, 16):
        field = f'scope_3_category_{i}_emissions'
        value = context.get(field, 0)
        if not value or value <= 0:
            categories_not_in_scope.append(i)

    # 生成汇总说明文字
    if categories_not_in_scope:
        category_str = "、".join(f"类别{cat}" for cat in categories_not_in_scope)
        formatted_context['scope_3_categories_not_in_scope_summary'] = (
            f"范围三{category_str}产生的排放为0，排放量不在本次盘查范围内，本周期内不进行量化。"
        )
    else:
        formatted_context['scope_3_categories_not_in_scope_summary'] = ""

    # ========== 格式化活动数据汇总表（基于位置） ==========
    # 为 act_summary_loc 中的数值字段添加格式化处理
    if 'act_summary_loc' in context and context['act_summary_loc']:
        formatted_act_summary = []
        emission_fields_to_format = [
            'act_summary_loc',  # 模板期望的字段名（活动数据数值）
            'activity_data_location_based',  # 完整字段名（兼容）
            'CO2_emissions',
            'CH4_emissions',
            'N2O_emissions',
            'HFCs_emissions',
            'PFCs_emissions',
            'SF6_emissions',
            'NF3_emissions',
            'total_green_house_gas_emissions'
        ]

        # 初始化汇总值
        loc_sums = {
            'CO2_emissions': 0.0,
            'CH4_emissions': 0.0,
            'N2O_emissions': 0.0,
            'HFCs_emissions': 0.0,
            'PFCs_emissions': 0.0,
            'SF6_emissions': 0.0,
            'NF3_emissions': 0.0,
            'total_green_house_gas_emissions': 0.0
        }

        for item in context['act_summary_loc']:
            formatted_item = item.copy()

            # 为数值字段添加格式化版本
            for field in emission_fields_to_format:
                if field in item:
                    original_value = item[field]
                    # 尝试转换为浮点数进行汇总
                    try:
                        num_value = float(str(original_value).replace(',', '').replace(' ', '')) if original_value else 0
                        if field in loc_sums:
                            loc_sums[field] += num_value
                    except (ValueError, TypeError):
                        pass

                    # 添加格式化版本（非空检查：确保输出 0.00 而不是空字符串）
                    if field == 'act_summary_loc':
                        # 活动数据使用格式化版本，同时保留原始值
                        formatted_item[field] = format_number(original_value) if original_value else '0.00'
                    elif field == 'activity_data_location_based':
                        formatted_item[f'{field}_formatted'] = format_number(original_value) if original_value else '0.00'
                    else:
                        # 排放量字段
                        formatted_item[field] = format_number(original_value) if original_value else '0.00'

            formatted_act_summary.append(formatted_item)

        formatted_context['act_summary_loc'] = formatted_act_summary
        print(f"[活动数据汇总表] 已格式化 {len(formatted_act_summary)} 行数据")

        # 添加汇总行数据（模板期望的格式）
        formatted_context['loc_CO2_emissions_sum_formatted'] = format_number(loc_sums['CO2_emissions'])
        formatted_context['loc_CH4_emissions_sum_formatted'] = format_number(loc_sums['CH4_emissions'])
        formatted_context['loc_N2O_emissions_sum_formatted'] = format_number(loc_sums['N2O_emissions'])
        formatted_context['loc_HFCs_emissions_sum_formatted'] = format_number(loc_sums['HFCs_emissions'])
        formatted_context['loc_PFCs_emissions_sum_formatted'] = format_number(loc_sums['PFCs_emissions'])
        formatted_context['loc_SF6_emissions_sum_formatted'] = format_number(loc_sums['SF6_emissions'])
        formatted_context['loc_NF3_emissions_sum_formatted'] = format_number(loc_sums['NF3_emissions'])
        formatted_context['loc_total_green_house_gas_emissions_sum_formatted'] = format_number(loc_sums['total_green_house_gas_emissions'])
        print(f"[活动数据汇总表] 汇总行计算完成")
    else:
        formatted_context['act_summary_loc'] = []
        # 设置空的汇总值
        formatted_context['loc_CO2_emissions_sum_formatted'] = '0.00'
        formatted_context['loc_CH4_emissions_sum_formatted'] = '0.00'
        formatted_context['loc_N2O_emissions_sum_formatted'] = '0.00'
        formatted_context['loc_HFCs_emissions_sum_formatted'] = '0.00'
        formatted_context['loc_PFCs_emissions_sum_formatted'] = '0.00'
        formatted_context['loc_SF6_emissions_sum_formatted'] = '0.00'
        formatted_context['loc_NF3_emissions_sum_formatted'] = '0.00'
        formatted_context['loc_total_green_house_gas_emissions_sum_formatted'] = '0.00'

    # ========== 格式化活动数据汇总表（基于市场） ==========
    # 为 act_summary_mar 中的数值字段添加格式化处理
    if 'act_summary_mar' in context and context['act_summary_mar']:
        formatted_act_summary_mar = []
        emission_fields_to_format_mar = [
            'act_summary_mar',  # 模板期望的字段名（活动数据数值）
            'activity_data_market_based',  # 完整字段名（兼容）
            'CO2_emissions',
            'CH4_emissions',
            'N2O_emissions',
            'HFCs_emissions',
            'PFCs_emissions',
            'SF6_emissions',
            'NF3_emissions',
            'total_green_house_gas_emissions'
        ]

        # 初始化汇总值
        mar_sums = {
            'CO2_emissions': 0.0,
            'CH4_emissions': 0.0,
            'N2O_emissions': 0.0,
            'HFCs_emissions': 0.0,
            'PFCs_emissions': 0.0,
            'SF6_emissions': 0.0,
            'NF3_emissions': 0.0,
            'total_green_house_gas_emissions': 0.0
        }

        for item in context['act_summary_mar']:
            formatted_item = item.copy()

            # 为数值字段添加格式化版本
            for field in emission_fields_to_format_mar:
                if field in item:
                    original_value = item[field]
                    # 尝试转换为浮点数进行汇总
                    try:
                        num_value = float(str(original_value).replace(',', '').replace(' ', '')) if original_value else 0
                        if field in mar_sums:
                            mar_sums[field] += num_value
                    except (ValueError, TypeError):
                        pass

                    # 添加格式化版本（非空检查：确保输出 0.00 而不是空字符串）
                    if field == 'act_summary_mar':
                        # 活动数据使用格式化版本，同时保留原始值
                        formatted_item[field] = format_number(original_value) if original_value else '0.00'
                    elif field == 'activity_data_market_based':
                        formatted_item[f'{field}_formatted'] = format_number(original_value) if original_value else '0.00'
                    else:
                        # 排放量字段
                        formatted_item[field] = format_number(original_value) if original_value else '0.00'

            formatted_act_summary_mar.append(formatted_item)

        formatted_context['act_summary_mar'] = formatted_act_summary_mar
        print(f"[活动数据汇总表] 已格式化 {len(formatted_act_summary_mar)} 行数据（基于市场）")

        # 添加汇总行数据（模板期望的格式）
        formatted_context['mar_CO2_emissions_sum_formatted'] = format_number(mar_sums['CO2_emissions'])
        formatted_context['mar_CH4_emissions_sum_formatted'] = format_number(mar_sums['CH4_emissions'])
        formatted_context['mar_N2O_emissions_sum_formatted'] = format_number(mar_sums['N2O_emissions'])
        formatted_context['mar_HFCs_emissions_sum_formatted'] = format_number(mar_sums['HFCs_emissions'])
        formatted_context['mar_PFCs_emissions_sum_formatted'] = format_number(mar_sums['PFCs_emissions'])
        formatted_context['mar_SF6_emissions_sum_formatted'] = format_number(mar_sums['SF6_emissions'])
        formatted_context['mar_NF3_emissions_sum_formatted'] = format_number(mar_sums['NF3_emissions'])
        formatted_context['mar_total_green_house_gas_emissions_sum_formatted'] = format_number(mar_sums['total_green_house_gas_emissions'])
        print(f"[活动数据汇总表] 汇总行计算完成（基于市场）")
    else:
        formatted_context['act_summary_mar'] = []
        # 设置空的汇总值
        formatted_context['mar_CO2_emissions_sum_formatted'] = '0.00'
        formatted_context['mar_CH4_emissions_sum_formatted'] = '0.00'
        formatted_context['mar_N2O_emissions_sum_formatted'] = '0.00'
        formatted_context['mar_HFCs_emissions_sum_formatted'] = '0.00'
        formatted_context['mar_PFCs_emissions_sum_formatted'] = '0.00'
        formatted_context['mar_SF6_emissions_sum_formatted'] = '0.00'
        formatted_context['mar_NF3_emissions_sum_formatted'] = '0.00'
        formatted_context['mar_total_green_house_gas_emissions_sum_formatted'] = '0.00'

    # ========== 格式化排放因子汇总表 ==========
    # 为 pro_ef_items 中的数值字段添加格式化处理
    if 'pro_ef_items' in context and context['pro_ef_items']:
        formatted_pro_ef_items = []
        ef_fields_to_format = [
            'ncv',                  # 低位发热量
            'ox_rate',              # 氧化率
            'ef_val',               # 计算值/排放系数
            'CO2_emission_factor',  # CO2排放因子
            'CH4_emission_factor',  # CH4排放因子
            'N2O_emission_factor',  # N2O排放因子
        ]

        # 初始化汇总值（排放因子汇总表通常需要计算平均值的汇总）
        ef_sums = {
            'ncv': 0.0,
            'ox_rate': 0.0,
            'ef_val': 0.0,
            'CO2_emission_factor': 0.0,
            'CH4_emission_factor': 0.0,
            'N2O_emission_factor': 0.0,
        }
        ef_count = 0  # 有效数据行计数

        for item in context['pro_ef_items']:
            formatted_item = item.copy()
            is_valid_row = False  # 标记是否为有效数据行

            # 为数值字段添加格式化版本
            for field in ef_fields_to_format:
                if field in item:
                    original_value = item[field]
                    # 尝试转换为浮点数进行汇总
                    try:
                        num_value = float(str(original_value).replace(',', '').replace(' ', '')) if original_value else 0
                        if field in ef_sums:
                            ef_sums[field] += num_value
                            is_valid_row = True
                    except (ValueError, TypeError):
                        pass

                    # 添加格式化版本（非空检查：确保输出 0.00 而不是空字符串）
                    formatted_item[field] = format_number(original_value) if original_value else '0.00'

            if is_valid_row:
                ef_count += 1
            formatted_pro_ef_items.append(formatted_item)

        formatted_context['pro_ef_items'] = formatted_pro_ef_items
        # 同步到旧变量名（向后兼容）
        formatted_context['emission_factor_items'] = formatted_pro_ef_items
        print(f"[排放因子汇总表] 已格式化 {len(formatted_pro_ef_items)} 行数据")

        # 添加汇总行数据（排放因子汇总表通常显示平均值）
        if ef_count > 0:
            formatted_context['ef_ncv_sum_formatted'] = format_number(ef_sums['ncv'] / ef_count)
            formatted_context['ef_ox_rate_sum_formatted'] = format_number(ef_sums['ox_rate'] / ef_count)
            formatted_context['ef_ef_val_sum_formatted'] = format_number(ef_sums['ef_val'] / ef_count)
            formatted_context['ef_CO2_emission_factor_sum_formatted'] = format_number(ef_sums['CO2_emission_factor'] / ef_count)
            formatted_context['ef_CH4_emission_factor_sum_formatted'] = format_number(ef_sums['CH4_emission_factor'] / ef_count)
            formatted_context['ef_N2O_emission_factor_sum_formatted'] = format_number(ef_sums['N2O_emission_factor'] / ef_count)
        else:
            formatted_context['ef_ncv_sum_formatted'] = '0.00'
            formatted_context['ef_ox_rate_sum_formatted'] = '0.00'
            formatted_context['ef_ef_val_sum_formatted'] = '0.00'
            formatted_context['ef_CO2_emission_factor_sum_formatted'] = '0.00'
            formatted_context['ef_CH4_emission_factor_sum_formatted'] = '0.00'
            formatted_context['ef_N2O_emission_factor_sum_formatted'] = '0.00'
        print(f"[排放因子汇总表] 汇总行计算完成（基于{ef_count}行有效数据）")
    else:
        formatted_context['pro_ef_items'] = []
        formatted_context['emission_factor_items'] = []
        # 设置空的汇总值
        formatted_context['ef_ncv_sum_formatted'] = '0.00'
        formatted_context['ef_ox_rate_sum_formatted'] = '0.00'
        formatted_context['ef_ef_val_sum_formatted'] = '0.00'
        formatted_context['ef_CO2_emission_factor_sum_formatted'] = '0.00'
        formatted_context['ef_CH4_emission_factor_sum_formatted'] = '0.00'
        formatted_context['ef_N2O_emission_factor_sum_formatted'] = '0.00'

    # ========== 格式化范围一直接排放源清册数据 ==========
    # 为 scope1 排放项中的数值字段添加格式化处理（2位小数 + 千分位符）
    # 确保零值显示为 "0.00" 而不是空字符串
    scope1_emission_vars = [
        'scope1_stationary_combustion_emissions_items',  # 固定燃烧
        'scope1_mobile_combustion_emissions_items',      # 移动燃烧
        'scope1_fugitive_emissions_items',               # 逸散排放
        'scope1_process_emissions_items',                # 制程排放
    ]

    # 需要格式化的排放量字段
    scope1_emission_fields = [
        'CO2_emissions',
        'CH4_emissions',
        'N2O_emissions',
        'HFCs_emissions',
        'PFCs_emissions',
        'SF6_emissions',
        'NF3_emissions',
        'total_green_house_gas_emissions'
    ]

    # 列名映射：数据源使用SF6，模板使用SFs
    # 注意：需要同时输出 SF6_emissions (数据源) 和 SFs_emissions (模板)
    sf6_to_sfs_mapping = {
        'SF6_emissions': 'SFs_emissions'
    }

    for var_name in scope1_emission_vars:
        if var_name in context and context[var_name]:
            formatted_items = []
            for item in context[var_name]:
                formatted_item = item.copy()

                # 为每个排放量字段应用格式化
                for field in scope1_emission_fields:
                    if field in item:
                        original_value = item[field]
                        # 应用格式化：保留2位小数，添加千分位符
                        # format_number(0) 返回 "0.00"，format_number(1234567.89) 返回 "1,234,567.89"
                        formatted_value = format_number(original_value) if original_value is not None else '0.00'
                        formatted_item[field] = formatted_value

                        # 特殊处理：SF6_emissions 同时映射为 SFs_emissions（模板兼容）
                        if field == 'SF6_emissions':
                            formatted_item['SFs_emissions'] = formatted_value

                formatted_items.append(formatted_item)

            formatted_context[var_name] = formatted_items
            print(f"[范围一排放] 已格式化 {var_name}: {len(formatted_items)} 行数据")
        else:
            # 如果变量不存在或为空，初始化为空列表
            formatted_context[var_name] = []

    # ========== 范围一排放格式化结束 ==========

    # 添加中文数字映射（用于自动编号）
    formatted_context['cn_nums'] = {
        1: '一', 2: '二', 3: '三', 4: '四', 5: '五',
        6: '六', 7: '七', 8: '八', 9: '九', 10: '十',
        11: '十一', 12: '十二', 13: '十三', 14: '十四', 15: '十五',
        16: '十六', 17: '十七', 18: '十八', 19: '十九', 20: '二十',
        21: '二十一', 22: '二十二', 23: '二十三', 24: '二十四', 25: '二十五',
        26: '二十六', 27: '二十七', 28: '二十八', 29: '二十九', 30: '三十'
    }
    
    # ========== 新增：最终字符串清洗步骤 ==========
    # 确保所有字符串值都已strip()，去除首尾空格
    def final_string_clean(d):
        """递归清洗所有字符串值"""
        if isinstance(d, dict):
            return {k: final_string_clean(v) for k, v in d.items()}
        elif isinstance(d, list):
            return [final_string_clean(item) for item in d]
        elif isinstance(d, str):
            return d.strip()
        else:
            return d

    formatted_context = final_string_clean(formatted_context)
    # ========== 最终清洗步骤结束 ==========

    # ========== 计算总温室气体排放量（基于位置和基于市场） ==========
    def safe_float(value):
        """安全转换为浮点数"""
        try:
            return float(value) if value is not None else 0.0
        except (ValueError, TypeError):
            return 0.0

    # 基于位置的总排放量 = 范围一 + 范围二（基于位置） + 范围三
    location_based_total = (
        safe_float(context.get('scope_1_emissions', 0)) +
        safe_float(context.get('scope_2_location_based_emissions', 0)) +
        safe_float(context.get('scope_3_emissions', 0))
    )
    formatted_location_total = format_number(location_based_total)
    formatted_context['location_based_total_green_house_gas_emissions'] = formatted_location_total
    formatted_context['location_based_total_green_house_gas_emissions_formatted'] = formatted_location_total

    # 基于市场的总排放量 = 范围一 + 范围二（基于市场） + 范围三
    market_based_total = (
        safe_float(context.get('scope_1_emissions', 0)) +
        safe_float(context.get('scope_2_market_based_emissions', 0)) +
        safe_float(context.get('scope_3_emissions', 0))
    )
    formatted_market_total = format_number(market_based_total)
    formatted_context['market_based_total_green_house_gas_emissions'] = formatted_market_total
    formatted_context['market_based_total_green_house_gas_emissions_formatted'] = formatted_market_total

    # Debug output
    import sys
    print(f"[总排放量计算]", file=sys.stderr)
    print(f"  基于位置的总排放量: {format_number(location_based_total)} 吨CO2e", file=sys.stderr)
    print(f"  基于市场的总排放量: {format_number(market_based_total)} 吨CO2e", file=sys.stderr)
    # ========== 总排放量计算结束 ==========

    return formatted_context


def generate_report_from_xlsx(
    xlsx_path="test_data.xlsx",
    template_path="template.docx",
    output_path="carbon_report.docx"
):
    """
    使用 template.docx 作为模板，从 xlsx 文件动态读取数据生成报告
    数字格式化在展示层通过 Jinja2 过滤器处理，数据层保持原始数字类型

    Args:
        xlsx_path: Excel 数据文件路径
        template_path: Word 模板文件路径
        output_path: 输出报告路径（默认: carbon_report.docx）
    """
    print("=" * 50)
    print("开始生成碳盘查报告（纯xlsx，动态读取）")
    print("=" * 50)

    # 1. 使用 data_reader 的协议驱动方法（重构版）
    print(f"\n[步骤1] 从 {xlsx_path} 动态提取数据...")
    reader = ExcelDataReader(xlsx_path)
    context = reader.get_all_context()  # 重构后使用 get_all_context()
    reader.close()  # 重构后需要手动关闭工作簿

    # 打印提取的关键数据
    print("\n提取的关键数据:")
    print(f"  公司名称: {context.get('company_name')}")
    print(f"  范围一排放: {context.get('scope_1_emissions')}")
    print(f"  范围二排放（基于位置）: {context.get('scope_2_location_based_emissions')}")
    print(f"  范围二排放（基于市场）: {context.get('scope_2_market_based_emissions')}")
    print(f"  范围三排放: {context.get('scope_3_emissions')}")

    print("\n范围三分类排放量:")
    for i in range(1, 16):  # 范围三共15个类别
        val = context.get(f'scope_3_category_{i}_emissions')
        if val and val > 0:
            print(f"  类别{i}: {val}")

    # 2. 加载模板
    print(f"\n[步骤2] 加载模板: {template_path}")
    template = DocxTemplate(template_path)

    # 3. 准备展示层数据（格式化数字，保持数据层纯净）
    print("\n[步骤3] 准备展示层数据（格式化数字）...")
    render_context = prepare_context_with_formatting(context)

    # 4. 渲染模板
    print("[步骤4] 渲染模板...")
    
    # 注册自定义 Jinja2 过滤器
    from jinja2 import Environment
    from jinja2_filters import format_number, format_emission, register_filters_to_template

    env = Environment()
    env.filters['cn_num'] = to_chinese_num
    env.filters['format_number'] = format_number
    env.filters['format_emission'] = format_emission
    template.jinja_env = env

    print("[渲染] 已注册过滤器: cn_num, format_number, format_emission")
    
    template.render(render_context)

    # 5. 保存报告
    print(f"\n[步骤5] 保存报告到: {output_path}")
    template.save(output_path)

    # 5.3. 检查类别12排放因子表是否被渲染
    print(f"\n[步骤5.3] 检查类别12排放因子表...")
    from docx import Document
    doc_check = Document(output_path)
    cat12_ef_found = False
    cat11_ef_idx = None

    # 首先找到类别11的表格
    for i, table in enumerate(doc_check.tables):
        for row in table.rows[:3]:
            for cell in row.cells:
                if '类别11' in cell.text and '排放因子' in cell.text:
                    cat11_ef_idx = i
                    print(f"  找到类别11排放因子表: 表格{i}")
                    break
            if cat11_ef_idx is not None:
                break
        if cat11_ef_idx is not None:
            break

    # 检查类别11之后的表格
    if cat11_ef_idx is not None:
        print(f"  检查类别11之后的表格（表格{cat11_ef_idx+1}到{min(cat11_ef_idx+3, len(doc_check.tables))}）")
        for i in range(cat11_ef_idx+1, min(cat11_ef_idx+3, len(doc_check.tables))):
            table = doc_check.tables[i]
            first_cell = table.rows[0].cells[0].text[:40] if table.rows[0].cells else ''
            row_count = len(table.rows)
            print(f"    表格{i}: {row_count}行, 首单元格=\"{first_cell}\"")

    # 搜索类别12
    for i, table in enumerate(doc_check.tables):
        for row in table.rows[:3]:
            for cell in row.cells:
                if '类别12' in cell.text and '排放因子' in cell.text:
                    cat12_ef_found = True
                    print(f"  找到类别12排放因子表: 表格{i}")
                    break
            if cat12_ef_found:
                break
        if cat12_ef_found:
            break
    if not cat12_ef_found:
        print(f"  警告: 类别12排放因子表未被渲染!")
        print(f"  文档中共有 {len(doc_check.tables)} 个表格")

    # 5.5. 检查模板渲染后的数据（调试用）
    print(f"\n[步骤5.5] 检查模板渲染后的数据...")
    check_template_rendering(output_path)

    # 6. 统一公司简介和经营范围的段落格式
    print(f"\n[步骤6] 统一段落格式...")
    from docx import Document
    from docx.shared import Pt, Inches, Cm

    doc = Document(output_path)
    company_name = context.get('company_name', '')

    # 设置统一的首行缩进：4个空格 ≈ 2个中文字符 ≈ 0.4厘米
    first_line_indent = Cm(0.4)  # 约2个中文字符的宽度

    # 遍历所有段落，精确处理公司简介和经营范围
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue

        # 检查是否是公司简介段落（以公司名开头或包含公司简介特征）
        if company_name in text[:50] or '楚能新能源' in text[:50]:
            if len(text) > 100:  # 确保是长文本内容而非标题
                # 设置统一的首行缩进，左缩进为0
                para.paragraph_format.first_line_indent = first_line_indent
                para.paragraph_format.left_indent = 0
                print(f"  已处理公司简介段落（长度: {len(text)} 字符）")

        # 检查是否是经营范围段落（包含经营范围特征）
        elif '经营范围' in text and len(text) > 100:
            # 设置统一的首行缩进，左缩进为0
            para.paragraph_format.first_line_indent = first_line_indent
            para.paragraph_format.left_indent = 0
            print(f"  已处理经营范围段落（长度: {len(text)} 字符）")

    # 7. 设置表2的列宽相等
    print(f"\n[步骤7] 设置表2列宽...")
    # 找到表2（范围二、三间接排放源表格）
    if len(doc.tables) >= 2:
        table2 = doc.tables[1]  # 第二个表格是表2
        # 设置所有列宽相等（使用 Inches 单位）
        equal_width = Inches(2.0)  # 每列2英寸
        for row in table2.rows:
            for cell in row.cells:
                # 设置单元格宽度
                cell.width = equal_width
        print("  表2列宽已设置为相等")

    doc.save(output_path)
    print("段落格式统一完成")

    # 8. 清理量化方法说明部分的过多空行
    print(f"\n[步骤8] 清理量化方法说明部分的空行...")
    clean_excessive_blank_lines(doc)

    doc.save(output_path)
    print("空行清理完成")

    # 9. 删除没有数据的类别表格（仅删除标题段落，保留表格结构）
    print(f"\n[步骤9] 删除没有数据的类别表格...")
    clean_empty_category_tables_v2(doc, context)

    doc.save(output_path)
    print("空类别表格清理完成")

    # 9.4. 修复范围三类别标题缺失类别名称的问题
    print(f"\n[步骤9.4] 修复范围三类别标题...")
    fix_scope3_category_headers(doc)

    doc.save(output_path)
    print("范围三类别标题修复完成")

    # 9.5. 检查合并前的表格数据
    print(f"\n[步骤9.5] 检查合并前的表格数据...")
    check_table_before_merge(output_path)

    # 10. 使用 XML vMerge 方法合并表格中的纵向单元格（针对表1和表2）
    print(f"\n[步骤10] 使用 XML vMerge 方法合并表格中的纵向单元格...")

    # 处理表1（doc.tables[0]）
    if len(doc.tables) >= 1:
        table1 = doc.tables[0]
        print(f"  处理表1（表格索引0）...")
        try:
            merge_vertical_cells(table1, 0)
        except Exception as e:
            print(f"  处理表1时出错: {e}")
            import traceback
            traceback.print_exc()

    # 处理表2（doc.tables[1]）
    if len(doc.tables) >= 2:
        table2 = doc.tables[1]
        print(f"  处理表2（表格索引1）...")
        try:
            merge_vertical_cells(table2, 0)
        except Exception as e:
            print(f"  处理表2时出错: {e}")
            import traceback
            traceback.print_exc()

    doc.save(output_path)
    print("表格纵向单元格合并完成（XML vMerge 方法）")

    # 10.5. 合并其他表格中的纵向单元格（XML方法）- 范围三类别表格
    print(f"\n[步骤10.5] 合并范围三类别表格的纵向单元格（XML方法）...")
    merge_other_tables_vertical_cells(doc, context)

    doc.save(output_path)
    print("范围三类别表格纵向单元格合并完成（XML方法）")

    print("\n" + "=" * 50)
    print(f"报告生成成功: {output_path}")
    print("=" * 50)

    return output_path


def check_template_rendering(doc_path):
    """
    检查模板渲染后的数据填充情况
    """
    from docx import Document

    try:
        doc = Document(doc_path)

        # 检查表格0
        if len(doc.tables) > 0:
            table = doc.tables[0]
            print(f"  表格0大小: {len(table.rows)} 行 x {len(table.columns)} 列")

            # 显示前5行的类别列数据
            print(f"  表格0第一列前5行:")
            for row_idx in range(min(5, len(table.rows))):
                cell_text = table.rows[row_idx].cells[0].text.strip()
                print(f"    第{row_idx}行: '{cell_text}'")

        # 检查 scope1 排放数据
        print(f"  scope1_stationary_combustion_emissions_items 前3条数据:")
        reader = ExcelDataReader('test_data.xlsx')
        data = reader.get_all_context()
        scope1_items = data.get('scope1_stationary_combustion_emissions_items', [])
        for i, item in enumerate(scope1_items[:3]):
            print(f"    {i+1}. category='{item.get('category')}', emission_source='{item.get('emission_source')}'")
        reader.close()

    except Exception as e:
        print(f"  检查时出错: {e}")


def check_table_before_merge(doc_path):
    """
    检查合并前的表格数据
    """
    from docx import Document

    try:
        doc = Document(doc_path)

        if len(doc.tables) > 0:
            table = doc.tables[0]
            print(f"  表格0大小: {len(table.rows)} 行 x {len(table.columns)} 列")

            # 检查前3行的数据
            print(f"  前3行数据:")
            for row_idx in range(min(3, len(table.rows))):
                col0_text = table.rows[row_idx].cells[0].text.strip()
                col1_text = table.rows[row_idx].cells[1].text.strip()
                print(f"    第{row_idx}行: 列0='{col0_text}', 列1='{col1_text}'")

    except Exception as e:
        print(f"  检查时出错: {e}")


def find_table_by_content(doc, search_keywords):
    """
    根据表格内容动态查找表格索引

    Args:
        doc: Word文档对象
        search_keywords: 要搜索的关键词列表（表格中包含任一关键词即匹配）

    Returns:
        匹配的表格索引，如果未找到返回None
    """
    for idx, table in enumerate(doc.tables):
        # 获取表格中所有文本
        table_text = ""
        try:
            # 直接访问XML元素，避免python-docx的导航问题
            for row in table._element.tr_lst:
                for tc in row.tc_lst:
                    # 获取单元格文本
                    for p in tc.p_lst:
                        for r in p.r_lst:
                            for t in r.t_lst:
                                table_text += t.text + " "
        except Exception as e:
            # 如果遍历出错（如合并单元格），跳过该表格
            continue

        # 检查是否包含任一关键词
        for keyword in search_keywords:
            if keyword in table_text:
                return idx

    return None


def find_summary_table(doc):
    """
    查找范围三排放汇总表格

    Returns:
        汇总表格的索引，如果未找到返回None
    """
    for idx, table in enumerate(doc.tables):
        # 汇总表格的特征：包含多个"范围三 类别X"和"排放量(tCO2e)"等
        table_text = ""
        for row in table.rows:
            for cell in row.cells:
                table_text += cell.text + " "

        # 检查是否是汇总表格（包含至少5个类别，且有"排放量"列）
        category_count = table_text.count("范围三 类别")
        
        # Skip debug output
        
        if category_count >= 5 and "排放量" in table_text:
            return idx

    return None


def clean_excessive_blank_lines(doc):
    """
    清理量化方法说明部分过多的空行
    策略：确保排放源之间只有一行物理间距
    """
    print("  正在清理量化方法说明部分的空行...")

    # 找到量化方法说明章节
    start_idx = None
    for i, para in enumerate(doc.paragraphs):
        if '量化方法说明' in para.text:
            start_idx = i
            break

    if not start_idx:
        print("  未找到量化方法说明章节")
        return

    # 找到该部分的结束位置
    end_idx = start_idx + 1
    for i in range(start_idx + 1, len(doc.paragraphs)):
        text = doc.paragraphs[i].text.strip()
        # 找到下一个主要章节
        if text and ('四、' in text or '参考文献' in text or '附录' in text):
            end_idx = i
            break
        if i > start_idx + 200:
            end_idx = i
            break

    print(f"  量化方法说明部分: {start_idx} 到 {end_idx}")

    # 在该部分内，删除连续的空段落，保留最多1个空行
    consecutive_empty = 0
    indices_to_remove = []

    for i in range(start_idx, end_idx):
        if i >= len(doc.paragraphs):
            break
        text = doc.paragraphs[i].text.strip()
        if not text:
            consecutive_empty += 1
            # 如果超过1个连续空行，标记删除
            if consecutive_empty > 1:
                indices_to_remove.append(i)
        else:
            consecutive_empty = 0

    # 从后往前删除
    removed_count = 0
    for idx in sorted(indices_to_remove, reverse=True):
        if idx < len(doc.paragraphs):
            para = doc.paragraphs[idx]
            parent = para._element.getparent()
            parent.remove(para._element)
            removed_count += 1

    print(f"  删除了 {removed_count} 个多余空行")




def clean_empty_category_tables(doc, context):
    """
    彻底删除没有数据的类别（类别8、13-14、15）的所有内容

    删除策略：
    1. 查找空类别的所有相关段落和表格
    2. 删除标题段落
    3. 删除紧邻的"单位：吨CO2e"段落（无论是在标题前面还是后面）
    4. 删除空表格
    5. 扫描整个文档，删除遗留的孤立"单位：吨CO2e"段落
    """
    # 范围三类别数量常量
    TOTAL_SCOPE3_CATEGORIES = 15

    # 类别编号到名称的映射
    category_names = {
        1: "购买的商品和服务",
        2: "资本商品",
        3: "燃料和能源相关活动",
        4: "上游运输和配送",
        5: "运营中产生的废弃物",
        6: "员工商务旅行",
        7: "员工通勤",
        8: "上游租赁资产",
        9: "下游运输和配送",
        10: "销售产品的加工",
        11: "销售产品的使用",
        12: "寿命终结处理",
        13: "下游租赁资产",
        14: "特许经营",
        15: "投资"
    }

    # 检查哪些类别没有数据
    # 判定条件：既没有排放因子数据，也没有详细数据，也没有排放量
    empty_categories = []
    for i in range(1, TOTAL_SCOPE3_CATEGORIES + 1):
        ef_items = context.get(f'cat{i}_ef_items', [])
        detail_items = context.get(f'scope3_category{i}', [])
        emission_value = context.get(f'scope_3_category_{i}_emissions', 0)

        has_ef_items = ef_items and len(ef_items) > 0
        has_detail_items = detail_items and len(detail_items) > 0
        has_emissions = emission_value and emission_value > 0

        # 如果没有任何数据，则标记为空类别
        if not (has_ef_items or has_detail_items or has_emissions):
            empty_categories.append(i)

    if not empty_categories:
        print("  所有类别都有数据，无需删除空类别表格")
        return

    print(f"  没有数据的类别: {empty_categories}")

    deleted_count = 0

    # 新的删除策略：先收集所有要删除的段落对象，再统一删除
    # 这样避免了删除过程中索引变化的问题
    all_paragraphs_to_remove = []

    # 步骤1：查找并删除所有相关的标题段落和单位段落
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue

        # 检查是否是空类别相关的段落
        is_empty_category_para = False
        matched_cat_num = None

        for cat_num in empty_categories:
            category_name = category_names.get(cat_num, "")
            is_target_category = (
                f'范围三 类别{cat_num}' in text or
                f'范围三类别{cat_num}' in text or
                f'类别{cat_num}' in text or
                (category_name and category_name in text)
            )

            if is_target_category:
                is_empty_category_para = True
                matched_cat_num = cat_num
                break

        if is_empty_category_para:
            all_paragraphs_to_remove.append(para)

    # 步骤2：查找与空类别关联的"单位：吨CO2e"段落
    # 通过检查段落的前后来判断
    all_paragraphs_list = list(doc.paragraphs)
    for i, para in enumerate(all_paragraphs_list):
        text = para.text.strip()
        if ('单位：吨CO2e' in text or '单位: 吨CO2e' in text) and para not in all_paragraphs_to_remove:
            # 检查这个单位段落的前后是否有空类别的标题
            check_range = 2  # 检查前后2个段落
            for j in range(max(0, i - check_range), min(len(all_paragraphs_list), i + check_range + 1)):
                if j == i:
                    continue
                nearby_text = all_paragraphs_list[j].text.strip()
                for cat_num in empty_categories:
                    if f'类别{cat_num}' in nearby_text:
                        all_paragraphs_to_remove.append(para)
                        break
                if para in all_paragraphs_to_remove:
                    break

    # 步骤3：删除所有标记的段落
    # 使用段落对象直接删除，避免索引问题
    deleted_count = 0
    for para in all_paragraphs_to_remove:
        try:
            para_element = para._element
            parent = para_element.getparent()
            if parent is not None:
                parent.remove(para_element)
                deleted_count += 1
        except Exception as e:
            print(f"  删除段落时出错: {e}")

    print(f"  已删除 {deleted_count} 个空类别相关段落")

    # 步骤4：删除空类别相关的表格
    # 表格26-40对应范围三类别1-15的排放因子表，需要根据空类别列表删除相应表格
    tables_to_remove = []

    # 范围三类别表格索引映射（模板中的固定位置）
    # 表格26-40分别对应类别1-15的排放因子表
    # 注意：类别11使用燃烧格式（表格36），类别12使用通用格式（表格37）
    scope3_ef_table_mapping = {
        1: 26,   # 表格26: 类别1
        2: 27,   # 表格27: 类别2
        3: 28,   # 表格28: 类别3
        4: 29,   # 表格29: 类别4
        5: 30,   # 表格30: 类别5
        6: 31,   # 表格31: 类别6
        7: 32,   # 表格32: 类别7
        8: 33,   # 表格33: 类别8
        9: 34,   # 表格34: 类别9
        10: 35,  # 表格35: 类别10
        11: 36,  # 表格36: 类别11 (燃烧格式)
        12: 37,  # 表格37: 类别12 (通用格式)
        13: 38,  # 表格38: 类别13
        14: 39,  # 表格39: 类别14
        15: 40,  # 表格40: 类别15
    }

    # 找出需要删除的表格索引（基于空排放因子表）
    for cat_num in empty_ef_table_categories:
        if cat_num in scope3_ef_table_mapping:
            table_idx = scope3_ef_table_mapping[cat_num]
            if table_idx < len(doc.tables):
                table = doc.tables[table_idx]
                # 检查表格是否只有表头（模板中默认5行，渲染后如果没有数据仍然是5行左右）
                # 检查是否有实际数据：查看第2行之后是否有非空内容
                has_data = False
                for row_idx in range(1, len(table.rows)):
                    row = table.rows[row_idx]
                    for cell in row.cells:
                        text = cell.text.strip()
                        # 排除表头关键词
                        header_keywords = ['编号', 'GHG排放类别', '排放源', '缺省排放因子',
                                        'CO2', 'CH4', 'N2O', 'HFCs', 'PFCs', 'SF6', 'NF3',
                                        '单位', '引用源', '基于热值']
                        if text and text not in header_keywords:
                            try:
                                val = float(text)
                                if val > 0:
                                    has_data = True
                                    break
                            except (ValueError, TypeError):
                                if len(text) > 3:  # 超过3个字符的文本
                                    has_data = True
                                    break
                    if has_data:
                        break

                if not has_data:
                    tables_to_remove.append(table_idx)
                    print(f"  标记删除类别{cat_num}的排放因子表: 索引{table_idx}")

    # 同时检查所有范围三类别表格（表格26-40），删除只有表头的表格
    print(f"  [后处理] 检查表格26-40（文档共有{len(doc.tables)}个表格）")
    for table_idx in range(26, min(41, len(doc.tables))):
        if table_idx in tables_to_remove:
            continue  # 已经标记删除

        table = doc.tables[table_idx]
        row_count = len(table.rows)

        # 调试：打印表格信息
        if row_count <= 6:
            first_cell = table.rows[0].cells[0].text[:30] if table.rows[0].cells else ''
            print(f"  [后处理] 表格{table_idx}: {row_count}行, 首单元格=\"{first_cell}\"")

        # 首先检查表格行数，如果只有2行（表头+子表头），直接判定为空表格
        if row_count <= 2:
            tables_to_remove.append(table_idx)
            print(f"  标记删除空表格: 索引{table_idx}（只有表头，行数={row_count}）")
            continue

        # 对于3-5行的表格，需要更仔细地检查是否有数据
        if row_count <= 5:
            # 定义表头关键词列表
            header_keywords = ['编号', 'GHG排放类别', '排放源', 'Activity name',
                             'Geography', 'CO2', 'CH4', 'N2O', '单位', '引用源',
                             '排放因子', '缺省', '基于热值']

            has_data = False
            # 从第2行开始检查（跳过前两行表头）
            for row_idx in range(2, min(row_count, 10)):
                row = table.rows[row_idx]

                # 检查第一列是否是数字编号（数据行的特征）
                if len(row.cells) > 0:
                    first_cell_text = row.cells[0].text.strip()
                    try:
                        num_val = float(first_cell_text)
                        if num_val > 0:  # 有编号，说明有数据行
                            has_data = True
                            break
                    except (ValueError, TypeError):
                        pass

            if not has_data:
                tables_to_remove.append(table_idx)
                print(f"  标记删除空表格: 索引{table_idx}（只有表头，行数={row_count}）")

    # 从后往前删除表格
    for table_idx in sorted(tables_to_remove, reverse=True):
        if table_idx < len(doc.tables):
            table = doc.tables[table_idx]
            table_element = table._element
            table_element.getparent().remove(table_element)
            deleted_count += 1

    # 步骤5：最终扫描 - 删除孤立的"单位：吨CO2e"段落
    # 检查所有"单位：吨CO2e"段落，如果它们不属于有数据的类别，则删除
    isolated_units = []
    for i, para in enumerate(doc.paragraphs):
        if '单位：吨CO2e' in para.text or '单位: 吨CO2e' in para.text:
            # 检查前后是否有有效的内容
            has_valid_content = False

            # 检查前5行和后5行
            start = max(0, i - 5)
            end = min(len(doc.paragraphs), i + 6)

            for j in range(start, end):
                if j == i:
                    continue
                text = doc.paragraphs[j].text.strip()
                # 检查是否有有效的类别内容
                for cat_num in range(1, 16):
                    if cat_num not in empty_categories:
                        if f'类别{cat_num}' in text or f'范围三 类别{cat_num}' in text:
                            has_valid_content = True
                            break
                if has_valid_content:
                    break

            if not has_valid_content:
                isolated_units.append(i)

    # 删除孤立的单位段落
    for idx in sorted(isolated_units, reverse=True):
        if idx < len(doc.paragraphs):
            para = doc.paragraphs[idx]
            para_element = para._element
            para_element.getparent().remove(para_element)
            deleted_count += 1

    print(f"  彻底删除完成，共删除 {deleted_count} 个元素")
    print(f"  包括标题段落、单位段落、表格和孤立单位段落")

    # 步骤6：删除孤立的"排放因子表"标题
    # 这些标题在删除表格后变成了孤立的
    orphan_headers_deleted = 0
    paragraphs_list = list(doc.paragraphs)

    # 获取有数据的类别列表
    valid_categories = []
    for i in range(1, 16):
        ef_items = context.get(f'cat{i}_ef_items', [])
        detail_items = context.get(f'scope3_category{i}', [])
        emission_value = context.get(f'scope_3_category_{i}_emissions', 0)
        if ef_items or detail_items or emission_value:
            valid_categories.append(i)

    for i, para in enumerate(paragraphs_list):
        text = para.text.strip()
        # 查找"范围三 类别X 排放因子表"标题
        import re
        match = re.search(r'范围三[^\d]*?(\d+)[^\d]*?排放因子表', text)
        if match:
            cat_num = int(match.group(1))
            # 检查这个类别是否是有效类别
            if cat_num not in valid_categories:
                # 这是一个空类别的标题，需要删除
                try:
                    para_element = para._element
                    parent = para_element.getparent()
                    if parent is not None:
                        parent.remove(para_element)
                        orphan_headers_deleted += 1
                        print(f"  删除孤立标题: 类别{cat_num}的排放因子表标题（无数据）")
                except Exception as e:
                    print(f"  删除孤立标题时出错: {e}")
            else:
                # 检查这个标题后面是否有表格（如果有数据，应该有表格）
                # 检查后面是否有表格元素，而不仅仅是文本内容
                has_table_after = False
                para_element = para._element
                current = para_element
                # 检查后面10个元素
                for _ in range(10):
                    current = current.getnext()
                    if current is None:
                        break
                    tag_name = current.tag.split('}')[-1] if '}' in current.tag else current.tag
                    if tag_name == 'tbl':
                        # 找到表格
                        has_table_after = True
                        break
                    # 如果遇到下一个主要章节，停止查找
                    elif tag_name == 'p':
                        para_text = ''
                        for t in current.iter():
                            if t.tag.split('}')[-1] == 't' and t.text:
                                para_text += t.text
                        if para_text.strip() and ('量化排除' in para_text or '不确定性' in para_text or '数据分级' in para_text):
                            # 遇到下一个章节，没有找到表格
                            break

                # 如果没有找到表格，说明表格未被渲染，标题应该删除
                if not has_table_after:
                    try:
                        para_element = para._element
                        parent = para_element.getparent()
                        if parent is not None:
                            parent.remove(para_element)
                            orphan_headers_deleted += 1
                            print(f"  删除孤立标题: 类别{cat_num}的排放因子表标题（未找到对应的表格）")
                    except Exception as e:
                        print(f"  删除孤立标题时出错: {e}")

    if orphan_headers_deleted > 0:
        print(f"  已删除 {orphan_headers_deleted} 个孤立的排放因子表标题")


def fix_scope3_category_headers(doc):
    """
    修复第四章量化说明中范围三类别标题缺失类别名称的问题

    修复内容：
    - 将 "（一）" 替换为 "（一）购买的商品和服务"
    - 将 "（二）" 替换为 "（二）资本货物"
    - 其他类别的类似标题
    """
    print("  正在修复范围三类别标题...")

    # 类别编号到名称的映射（与模板中的顺序对应）
    category_names = {
        '（一）': '购买的商品和服务',
        '（二）': '资本货物',
        '（三）': '燃料和能源相关活动',
        '（四）': '上游运输和配送',
        '（五）': '运营中产生的废弃物',
        '（六）': '员工商务旅行',
        '（七）': '员工通勤',
        '（八）': '上游租赁资产',
        '（九）': '下游运输和配送',
        '（十）': '销售产品的加工',
        '（十一）': '销售产品的使用',
        '（十二）': '寿命终结处理',
        '（十三）': '下游租赁资产',
        '（十四）': '特许经营',
        '（十五）': '投资'
    }

    # 找到"范围三：其他间接温室气体排放"的位置
    scope3_section_start = None
    paragraphs_list = list(doc.paragraphs)

    for i, para in enumerate(paragraphs_list):
        text = para.text.strip()
        if '范围三' in text and '其他间接温室气体排放' in text:
            scope3_section_start = i
            break

    if scope3_section_start is None:
        print("    未找到范围三章节，跳过修复")
        return

    print(f"    找到范围三章节在段落 {scope3_section_start}")

    # 新增：在范围三标题之前插入分组标题"（三）范围三其他类别相关排放"
    # 检查是否已经有这个分组标题
    has_group_header = False
    for i in range(max(0, scope3_section_start - 3), scope3_section_start):
        if i < len(paragraphs_list):
            text = paragraphs_list[i].text.strip()
            if '（三）' in text and ('范围三' in text or '其他类别' in text):
                has_group_header = True
                break

    if not has_group_header:
        # 在"范围三：其他间接温室气体排放"之前插入分组标题
        target_para = paragraphs_list[scope3_section_start]

        # 创建新段落
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn

        new_para = OxmlElement('w:p')
        pPr = OxmlElement('w:pPr')
        jc = OxmlElement('w:jc')
        jc.set(qn('w:val'), 'left')
        pPr.append(jc)
        new_para.append(pPr)

        r = OxmlElement('w:r')
        rPr = OxmlElement('w:rPr')
        # 设置加粗
        b = OxmlElement('w:b')
        rPr.append(b)
        # 设置字号
        sz = OxmlElement('w:sz')
        sz.set(qn('w:val'), '24')  # 12pt
        rPr.append(sz)
        r.append(rPr)

        t = OxmlElement('w:t')
        t.text = '（三）范围三其他类别相关排放'
        r.append(t)
        new_para.append(r)

        # 在目标段落前插入新段落
        target_para._element.addprevious(new_para)
        print(f"    已在范围三章节前插入分组标题 '（三）范围三其他类别相关排放'")

    # 在范围三章节内查找只有编号没有类别名称的标题
    fixed_count = 0
    for i in range(scope3_section_start + 1, len(paragraphs_list)):
        text = paragraphs_list[i].text.strip()

        # 检查是否是纯编号标题（只有"（X）"而没有其他内容）
        if text in category_names:
            # 检查下一段落是否是"（1）量化模型"等，确认这确实是一个类别标题
            if i + 1 < len(paragraphs_list):
                next_text = paragraphs_list[i + 1].text.strip()
                if next_text.startswith('（1') or next_text.startswith('（2'):
                    # 这是一个需要修复的类别标题
                    full_title = f"{text}{category_names[text]}"
                    para = paragraphs_list[i]

                    # 清除段落内容
                    for run in para.runs:
                        run.text = ""

                    # 添加新文本
                    if para.runs:
                        para.runs[0].text = full_title
                    else:
                        para.add_run(full_title)

                    fixed_count += 1
                    print(f"    修复段落{i}: '{text}' -> '{full_title}'")

        # 如果到达下一个范围章节，停止处理
        if text and ('第四章' in text or '参考文献' in text or '附录' in text):
            break

    print(f"    已修复 {fixed_count} 个范围三类别标题")


def insert_cat12_emission_factor_table(doc, after_table_idx, cat12_items):
    """
    手动插入类别12排放因子表及其小标题

    Args:
        doc: Word文档对象
        after_table_idx: 在此表格索引后插入新表格
        cat12_items: 类别12的排放因子数据
    """
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    # 获取插入位置的表格元素
    if after_table_idx >= len(doc.tables):
        print(f"  错误: 表格索引{after_table_idx}超出范围")
        return

    ref_table = doc.tables[after_table_idx]
    ref_table_element = ref_table._element
    parent = ref_table_element.getparent()

    # 步骤1：插入类别12的小标题段落
    subtitle_para = OxmlElement('w:p')
    subtitle_para.set(qn('w:rsidR'), '007D2F2A')
    subtitle_para.set(qn('w:rsidRPr'), '007D2F2A')
    subtitle_ppr = OxmlElement('w:pPr')
    subtitle_jc = OxmlElement('w:jc')
    subtitle_jc.set(qn('w:val'), 'left')
    subtitle_ppr.append(subtitle_jc)

    # 小标题文本
    subtitle_r = OxmlElement('w:r')
    subtitle_r.set(qn('w:rsidRPr'), '007D2F2A')
    subtitle_rpr = OxmlElement('w:rPr')
    subtitle_b = OxmlElement('w:b')
    subtitle_rpr.append(subtitle_b)
    subtitle_r.append(subtitle_rpr)

    subtitle_t = OxmlElement('w:t')
    subtitle_t.text = '范围三 类别12 售出产品报废产生的排放 排放因子表'
    subtitle_r.append(subtitle_t)
    subtitle_para.append(subtitle_r)
    subtitle_para.append(subtitle_ppr)

    # 在参考表格后插入小标题段落
    parent.insert(parent.index(ref_table_element) + 1, subtitle_para)
    print(f"  插入类别12小标题段落")

    # 步骤2：创建新表格
    new_tbl = OxmlElement('w:tbl')

    # 添加表格属性
    tbl_pr = _get_table_properties()
    new_tbl.append(tbl_pr)

    # 添加表格边框
    tbl_borders = OxmlElement('w:tblBorders')

    # 定义边框
    for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '4')
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), 'auto')
        tbl_borders.append(border)

    tbl_pr.append(tbl_borders)

    # 第1行：主表头（8列）
    tr1 = OxmlElement('w:tr')
    tr1.set(qn('w:rsidR'), '007D2F2A')

    headers_row1 = ['编号', 'GHG排放类别', '排放源', 'Activity name', 'Geography', 'CO2', 'CO2', '引用源']
    for header in headers_row1:
        tc = _create_table_cell(header, is_header=True)
        tr1.append(tc)
    new_tbl.append(tr1)

    # 第2行：子表头（8列，前5列空，后3列有内容）
    tr2 = OxmlElement('w:tr')
    tr2.set(qn('w:rsidR'), '007D2F2A')

    headers_row2 = ['', '', '', '', '', '排放因子', '单位', '缺省排放因子']
    for header in headers_row2:
        tc = _create_table_cell(header, is_header=True)
        tr2.append(tc)
    new_tbl.append(tr2)

    # 数据行
    for item in cat12_items:
        tr = OxmlElement('w:tr')
        tr.set(qn('w:rsidR'), '007D2F2A')

        cells = [
            str(item.get('number', '')),
            item.get('emission_source_type_cat12', ''),
            item.get('emission_source_cat12', ''),
            item.get('emission_name_cat12', ''),
            item.get('emission_geo_cat12', ''),
            str(item.get('cat12_emission_factor', '')),
            item.get('cat12_emission_unit', ''),
            item.get('cat12_emission_source', ''),
        ]

        for cell_text in cells:
            tc = _create_table_cell(cell_text, is_header=False)
            tr.append(tc)

        new_tbl.append(tr)

    # 在小标题段落后插入表格
    parent.insert(parent.index(subtitle_para) + 1, new_tbl)
    print(f"  成功插入类别12排放因子表（{len(cat12_items)}行数据，8列）")


def _get_table_properties():
    """获取表格属性XML"""
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    tbl_pr = OxmlElement('w:tblPr')
    tbl_pr.set(qn('w:tblStyle'), 'TableGrid')

    # 表格宽度
    tbl_w = OxmlElement('w:tblW')
    tbl_w.set(qn('w:w'), '0')
    tbl_w.set(qn('w:type'), 'auto')
    tbl_pr.append(tbl_w)

    # 添加表格网格定义（8列，每列等宽）
    tbl_grid = OxmlElement('w:tblGrid')
    for i in range(8):
        grid_col = OxmlElement('w:gridCol')
        grid_col.set(qn('w:w'), '1250')  # 每列宽度
        tbl_grid.append(grid_col)
    tbl_pr.append(tbl_grid)

    return tbl_pr


def _create_table_cell(text, is_header=False):
    """创建表格单元格"""
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    tc = OxmlElement('w:tc')
    tc_pr = OxmlElement('w:tcPr')
    tc_w = OxmlElement('w:tcW')
    tc_w.set(qn('w:w'), '882')
    tc_w.set(qn('w:type'), 'pct')
    tc_pr.append(tc_w)

    # 边框
    tc_borders = OxmlElement('w:tcBorders')
    for border_pos in ['top', 'left', 'bottom', 'right']:
        border = OxmlElement(f'w:{border_pos}')
        border.set(qn('w:val'), 'nil' if is_header else 'single')
        if not is_header and border_pos in ['top', 'left']:
            border.set(qn('w:val'), 'single')
            border.set(qn('w:sz'), '4')
            border.set(qn('w:space'), '0')
            border.set(qn('w:color'), 'auto')
        tc_borders.append(border)

    tc_pr.append(tc_borders)
    tc.append(tc_pr)

    # 段落
    p = OxmlElement('w:p')
    p.set(qn('w:rsidR'), '007D2F2A')
    p.set(qn('w:rsidRPr'), '007D2F2A')
    p_pr = OxmlElement('w:pPr')
    p_jc = OxmlElement('w:jc')
    p_jc.set(qn('w:val'), 'center' if is_header else 'left')
    p_pr.append(p_jc)
    p.append(p_pr)

    # 文本
    r = OxmlElement('w:r')
    r.set(qn('w:rsidRPr'), '007D2F2A')
    r_pr = OxmlElement('w:rPr')
    if is_header:
        # 表头使用加粗
        b = OxmlElement('w:b')
        r_pr.append(b)
    r.append(r_pr)

    t = OxmlElement('w:t')
    t.text = text
    r.append(t)
    p.append(r)

    tc.append(p)

    return tc


def clean_empty_category_tables_v2(doc, context):
    """
    删除没有数据的类别的标题段落和单位段落，但保留表格结构

    新版本：不删除表格本身，避免表格索引偏移问题
    只删除与空类别相关的标题段落和单位段落
    """
    TOTAL_SCOPE3_CATEGORIES = 15

    category_names = {
        1: "购买的商品和服务",
        2: "资本商品",
        3: "燃料和能源相关活动",
        4: "上游运输和配送",
        5: "运营中产生的废弃物",
        6: "员工商务旅行",
        7: "员工通勤",
        8: "上游租赁资产",
        9: "下游运输和配送",
        10: "销售产品的加工",
        11: "销售产品的使用",
        12: "售出产品的加工",
        13: "下游租赁资产",
        14: "特许经营",
        15: "投资"
    }

    # 检查哪些类别完全没有任何数据（既没有ef_items，也没有detail_items，也没有emissions）
    empty_categories = []
    # 同时检查哪些类别的排放因子表只有表头（没有数据行）
    empty_ef_table_categories = []

    for i in range(1, TOTAL_SCOPE3_CATEGORIES + 1):
        ef_items = context.get(f'cat{i}_ef_items', [])
        detail_items = context.get(f'scope3_category{i}', [])
        emission_value = context.get(f'scope_3_category_{i}_emissions', 0)

        has_ef_items = ef_items and len(ef_items) > 0
        has_detail_items = detail_items and len(detail_items) > 0
        has_emissions = emission_value and emission_value > 0

        # 如果没有任何数据，则标记为完全空类别
        if not (has_ef_items or has_detail_items or has_emissions):
            empty_categories.append(i)

        # 如果排放因子表没有数据（ef_items为空），标记为需要删除排放因子表标题
        if not has_ef_items:
            empty_ef_table_categories.append(i)

    print(f"  空排放因子表类别: {empty_ef_table_categories}")  # 调试输出

    # 合并两类需要删除的类别
    categories_to_remove = list(set(empty_categories + empty_ef_table_categories))

    if not categories_to_remove:
        print("  所有类别都有数据，无需删除空类别标题")
        return

    print(f"  完全空类别（无任何数据）: {empty_categories}")
    print(f"  排放因子表为空的类别: {[c for c in empty_ef_table_categories if c not in empty_categories]}")

    # 收集要删除的段落
    all_paragraphs_to_remove = []

    # 步骤1：查找并删除空类别和空排放因子表的标题段落
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue

        # 检查是否是空类别或空排放因子表相关的段落
        is_empty_category_para = False

        for cat_num in categories_to_remove:
            category_name = category_names.get(cat_num, "")
            is_target_category = (
                f'范围三 类别{cat_num}' in text or
                f'范围三类别{cat_num}' in text or
                f'类别{cat_num} ' in text or  # 类别X后面有空格，避免匹配"类别10"到"类别1"
                (category_name and category_name in text)
            )

            if is_target_category:
                is_empty_category_para = True
                break

        if is_empty_category_para:
            all_paragraphs_to_remove.append(para)

    # 步骤2：删除所有标记的段落
    deleted_count = 0
    for para in all_paragraphs_to_remove:
        try:
            para_element = para._element
            parent = para_element.getparent()
            if parent is not None:
                parent.remove(para_element)
                deleted_count += 1
        except Exception as e:
            print(f"  删除段落时出错: {e}")

    print(f"  已删除 {deleted_count} 个空类别相关段落")

    # 步骤3：删除空类别相关的表格
    tables_to_remove = []

    # 定义表头关键词列表
    header_keywords = ['编号', 'GHG排放类别', '排放源', 'Activity name',
                     'Geography', 'CO2', 'CH4', 'N2O', '单位', '引用源',
                     '排放因子', '缺省', '基于热值']

    # 动态查找每个类别对应的表格
    category_table_indices = {}
    for cat_num in range(1, 16):
        # 在所有表格中查找属于该类别的表格（从表格4开始，包含范围三类别汇总表）
        for table_idx in range(4, min(50, len(doc.tables))):  # 扩大搜索范围，包含表格4-18
            if table_idx in category_table_indices.values():
                continue  # 已经被其他类别占用

            table = doc.tables[table_idx]
            row_count = len(table.rows)

            # 检查表格是否属于该类别
            is_category_table = False
            for row_idx in range(min(6, row_count)):  # 从第0行开始检查所有行
                row = table.rows[row_idx]
                for cell in row.cells:
                    text = cell.text.strip()
                    if f'范围三 类别{cat_num}' in text or f'范围三类别{cat_num}' in text:
                        is_category_table = True
                        break
                if is_category_table:
                    break

            if is_category_table:
                category_table_indices[cat_num] = table_idx
                print(f"  [DEBUG] 类别{cat_num} -> 表格{table_idx}（行数={row_count}）")
                break

    # 检查哪些类别的表格需要删除（只删除完全没有数据的类别）
    for cat_num in empty_categories:
        if cat_num in category_table_indices:
            table_idx = category_table_indices[cat_num]
            if table_idx < len(doc.tables):
                table = doc.tables[table_idx]
                row_count = len(table.rows)
                tables_to_remove.append(table_idx)
                print(f"  标记删除类别{cat_num}的排放因子表: 索引{table_idx}（行数={row_count}）")

    # 同时检查所有范围三类别表格，删除明显的空表格
    # 扩大搜索范围，确保覆盖所有可能的表格
    for table_idx in range(26, min(len(doc.tables), len(doc.tables))):  # 搜索到文档末尾
        if table_idx in tables_to_remove or table_idx in category_table_indices.values():
            continue  # 已经处理过

        table = doc.tables[table_idx]
        row_count = len(table.rows)

        # 删除明显的空表格（行数<=3）
        if row_count <= 3:
            # 检查表格是否是范围三类别表格
            is_scope3_table = False
            for row in table.rows[:3]:  # 检查前3行
                for cell in row.cells:
                    text = cell.text.strip()
                    if 'GHG排放类别' in text or '范围三' in text:
                        is_scope3_table = True
                        break
                if is_scope3_table:
                    break

            if is_scope3_table:
                tables_to_remove.append(table_idx)
                print(f"  标记删除明显空表格: 索引{table_idx}（范围三表格，只有表头，行数={row_count}）")
            continue

        # 对于4-6行的表格，检查是否有数据行
        if 4 <= row_count <= 6:
            has_data = False
            for row_idx in range(2, min(row_count, 10)):  # 从第2行开始检查
                row = table.rows[row_idx]
                if len(row.cells) > 0:
                    first_cell_text = row.cells[0].text.strip()
                    try:
                        num_val = float(first_cell_text)
                        if num_val > 0:  # 有编号，说明有数据行
                            has_data = True
                            break
                    except (ValueError, TypeError):
                        pass

            if not has_data:
                tables_to_remove.append(table_idx)
                print(f"  标记删除空表格: 索引{table_idx}（无数据行，行数={row_count}）")

    # 从后往前删除表格
    deleted_table_count = 0
    for table_idx in sorted(tables_to_remove, reverse=True):
        if table_idx < len(doc.tables):
            table = doc.tables[table_idx]
            table_element = table._element
            table_element.getparent().remove(table_element)
            deleted_table_count += 1

    print(f"  已删除 {deleted_table_count} 个空类别表格")


def merge_vertical_cells(table, col_idx):
    """
    纵向合并表格指定列中内容相同的相邻单元格

    使用底层 XML 操作，正确设置 vMerge 属性来实现合并。
    直接访问 XML 元素，避免 python-docx 的导航问题。

    Args:
        table: python-docx 表格对象
        col_idx: 要处理的列索引（从0开始）

    Returns:
        合并的单元格数量
    """
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    if col_idx >= len(table.columns):
        print(f"  警告：列索引 {col_idx} 超出表格范围（表格共 {len(table.columns)} 列）")
        return 0

    merged_count = 0
    rows_count = len(table.rows)

    # 跳过表头行，从第1行开始处理（索引1）
    start_row = 1

    # 第一步：直接从 XML 读取所有行的文本内容
    print(f"  开始识别合并组（总行数：{rows_count}，从行{start_row}开始）...")

    # 获取表格的 XML 元素
    table_element = table._element
    tr_lst = table_element.tr_lst

    # 存储每行的文本内容
    row_texts = []

    for row_idx in range(len(tr_lst)):
        if row_idx < start_row:
            # 跳过表头行
            row_texts.append("")
            continue

        try:
            tr = tr_lst[row_idx]
            tc_lst = tr.tc_lst

            if col_idx >= len(tc_lst):
                print(f"  警告：行{row_idx}的列索引{col_idx}超出范围")
                row_texts.append("")
                continue

            tc = tc_lst[col_idx]

            # 获取单元格文本
            cell_text = ""
            for p in tc.p_lst:
                for r in r_lst if (r_lst := p.r_lst) else []:
                    for t in t_lst if (t_lst := r.t_lst) else []:
                        cell_text += t.text

            cell_text = cell_text.strip()
            row_texts.append(cell_text)
        except Exception as e:
            print(f"  警告：读取行{row_idx}时出错: {e}")
            row_texts.append("")

    # 第二步：识别需要合并的单元格组
    merge_groups = []
    group_start = start_row
    current_value = None

    for row_idx in range(start_row, rows_count):
        cell_text = row_texts[row_idx] if row_idx < len(row_texts) else ""

        # 初始化当前值（第一行）
        if current_value is None:
            current_value = cell_text
            group_start = row_idx
            text_preview = cell_text[:40] if cell_text else ""
            print(f"  行{row_idx}: 初始化，内容=\"{text_preview}\"")
            continue

        # 如果当前单元格内容与前一单元格相同，继续累积
        if cell_text == current_value and cell_text != "":
            continue

        # 内容不同或遇到空值，记录之前的合并组
        if row_idx - 1 > group_start:
            merge_groups.append((group_start, row_idx - 1))
            value_preview = current_value[:40] if current_value else ""
            print(f"  识别合并组: 行{group_start}-{row_idx-1} (共{row_idx-1-group_start}行), 内容=\"{value_preview}\"")

        # 开始新的合并组
        current_value = cell_text
        group_start = row_idx
        if cell_text:
            text_preview = cell_text[:40] if cell_text else ""
            print(f"  行{row_idx}: 新组开始，内容=\"{text_preview}\"")

    # 处理最后一组
    if rows_count - 1 > group_start:
        merge_groups.append((group_start, rows_count - 1))
        value_preview = current_value[:40] if current_value else ""
        print(f"  识别合并组: 行{group_start}-{rows_count-1} (共{rows_count-1-group_start}行), 内容=\"{value_preview}\"")

    print(f"  总共识别到 {len(merge_groups)} 个合并组")

    # 第三步：执行合并操作（直接操作 XML）
    print(f"  开始执行合并操作，共 {len(merge_groups)} 个合并组...")

    for group_start, group_end in merge_groups:
        if group_end <= group_start:
            print(f"  跳过无效合并组: 行{group_start}-{group_end}")
            continue

        print(f"  处理合并组: 行{group_start}-{group_end} (共{group_end-group_start}行)")

        try:
            # 获取顶部单元格的 XML 元素
            tr_start = tr_lst[group_start]
            tc_lst_start = tr_start.tc_lst

            if col_idx >= len(tc_lst_start):
                print(f"    警告：列索引{col_idx}超出范围")
                continue

            tc_start = tc_lst_start[col_idx]

            # 保存顶部单元格的文本内容
            merged_text = row_texts[group_start] if group_start < len(row_texts) else ""
            text_preview = merged_text[:40] if merged_text else ""
            print(f"    顶部单元格（行{group_start}）内容: \"{text_preview}\"")

            # 设置顶部单元格的 vMerge 属性为 "restart"
            tc_pr = tc_start.get_or_add_tcPr()
            for old_vmerge in tc_pr.findall(qn('w:vMerge')):
                tc_pr.remove(old_vmerge)
            v_merge = OxmlElement('w:vMerge')
            v_merge.set(qn('w:val'), 'restart')
            tc_pr.append(v_merge)

            # 验证vMerge属性是否设置成功
            verify_vm = tc_pr.find(qn('w:vMerge'))
            verify_val = verify_vm.get(qn('w:val')) if verify_vm is not None else 'not found'
            print(f"    设置行{group_start} vMerge=\"restart\" (验证: {verify_val})")

            # 处理其他单元格（设置为 "continue"）
            for row_idx in range(group_start + 1, group_end + 1):
                if row_idx >= len(tr_lst):
                    break

                try:
                    tr = tr_lst[row_idx]
                    tc_lst = tr.tc_lst

                    if col_idx >= len(tc_lst):
                        continue

                    tc = tc_lst[col_idx]

                    # 清空被合并单元格的内容
                    for p in tc.p_lst:
                        # 移除所有 run 元素
                        for r in p.r_lst[:]:
                            p.remove(r)

                    # 设置 vMerge 属性为 "continue"
                    merge_tc_pr = tc.get_or_add_tcPr()
                    for old_vmerge in merge_tc_pr.findall(qn('w:vMerge')):
                        merge_tc_pr.remove(old_vmerge)
                    continue_vmerge = OxmlElement('w:vMerge')
                    continue_vmerge.set(qn('w:val'), 'continue')
                    merge_tc_pr.append(continue_vmerge)

                    merged_count += 1
                except Exception as e:
                    print(f"    警告：处理行{row_idx}时出错: {e}")

            print(f"    设置行{group_start+1}-{group_end} vMerge=\"continue\" (共{group_end-group_start}行)")

            # 设置顶部单元格的内容
            # 清空现有内容
            for p in tc_start.p_lst[:]:
                # 移除所有段落（除了保留一个）
                if len(tc_start.p_lst) > 1:
                    tc_start.remove(p)

            # 确保至少有一个段落
            if len(tc_start.p_lst) == 0:
                from docx.oxml import parse_xml
                new_p = parse_xml(r'<w:p xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"/>')
                tc_start.append(new_p)

            # 设置内容到第一个段落
            p = tc_start.p_lst[0]
            if merged_text:
                # 清空现有 run
                for r in p.r_lst[:]:
                    p.remove(r)

                # 添加新的 run 和 text
                from docx.oxml import parse_xml
                new_r = parse_xml(r'<w:r xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"><w:t xml:space="preserve"></w:t></w:r>')
                t_element = new_r.find(qn('w:t'))
                if t_element is not None:
                    t_element.text = merged_text
                p.append(new_r)

                content_preview = merged_text[:40] if merged_text else ""
                print(f"    设置行{group_start}内容: \"{content_preview}\"")

                # 验证内容是否设置成功
                verify_text = ""
                for r in p.r_lst:
                    for t in t_lst if (t_lst := r.t_lst) else []:
                        verify_text += t.text
                verify_preview = verify_text[:40] if verify_text else ""
                print(f"    验证行{group_start}内容: \"{verify_preview}\"")
            else:
                print(f"    [警告] merged_text 为空，跳过内容设置")

            # 设置合并后单元格的对齐方式
            try:
                # 验证vMerge属性在设置对齐方式之前是否仍然存在
                tc_pr_check = tc_start.get_or_add_tcPr()
                vmerge_check = tc_pr_check.find(qn('w:vMerge'))
                vmerge_val_check = vmerge_check.get(qn('w:val')) if vmerge_check is not None else 'None'
                print(f"    设置对齐方式前vMerge验证: vMerge=\"{vmerge_val_check}\"")

                # 垂直居中对齐
                tc_pr = tc_start.get_or_add_tcPr()
                v_align = tc_pr.find(qn('w:vAlign'))
                if v_align is None:
                    v_align = OxmlElement('w:vAlign')
                    v_align.set(qn('w:val'), 'center')
                    tc_pr.append(v_align)

                # 段落居中对齐
                jc = OxmlElement('w:jc')
                jc.set(qn('w:val'), 'center')
                p_pr = p.get_or_add_pPr()
                p_pr.append(jc)

                # 验证vMerge属性在设置对齐方式后是否仍然存在
                vmerge_check2 = tc_pr.find(qn('w:vMerge'))
                vmerge_val_check2 = vmerge_check2.get(qn('w:val')) if vmerge_check2 is not None else 'None'
                print(f"    设置对齐方式后vMerge验证: vMerge=\"{vmerge_val_check2}\"")
            except Exception as e:
                print(f"    设置单元格对齐方式时出错: {e}")

        except Exception as e:
            print(f"    处理合并组时出错: {e}")
            import traceback
            traceback.print_exc()

    print(f"  第 {col_idx} 列纵向合并完成，合并了 {merged_count} 个单元格")
    return merged_count


def merge_table_cells(table, col_idx):
    """
    纵向合并相同内容的单元格并居中（使用 XML vMerge 方法）

    使用 XML vMerge 属性进行单元格合并，这种方法更可靠。
    在合并前先清除所有现有的 vMerge 属性，避免干扰。

    Args:
        table: python-docx 表格对象
        col_idx: 要处理的列索引（从0开始）

    Returns:
        合并的单元格数量
    """
    from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    rows = table.rows
    if len(rows) < 2:
        return 0

    merged_count = 0

    # 跳过表头行，从第1行开始（索引1）
    # 假设第0行是表头，不需要合并
    if len(rows) < 3:
        return 0

    # 第一步：清除所有现有的 vMerge 属性
    for row in rows:
        if col_idx < len(row.cells):
            cell = row.cells[col_idx]
            tc_pr = cell._element.get_or_add_tcPr()
            # 移除所有 vMerge 属性
            for old_vmerge in tc_pr.findall(qn('w:vMerge')):
                tc_pr.remove(old_vmerge)

    # 第二步：收集所有单元格的文本内容
    cell_texts = []
    for row in rows[1:]:  # 跳过表头
        if col_idx < len(row.cells):
            cell = row.cells[col_idx]
            text = cell.text.strip()
            cell_texts.append(text)

    # 第三步：使用组检测逻辑，识别需要合并的单元格组
    merge_groups = []
    start_idx = 0
    current_text = cell_texts[0] if cell_texts else ""

    for i in range(1, len(cell_texts)):
        if cell_texts[i] == current_text and current_text != "":
            # 继续累积，等待发现不同的内容
            continue
        else:
            # 内容不同或遇到空值，记录之前的组
            if i - 1 > start_idx:
                merge_groups.append((start_idx, i - 1))
            # 开始新的组
            current_text = cell_texts[i]
            start_idx = i

    # 处理最后一组
    if len(cell_texts) - 1 > start_idx:
        merge_groups.append((start_idx, len(cell_texts) - 1))

    # 调试输出
    print(f"  识别到 {len(merge_groups)} 个合并组:")
    for i, (group_start, group_end) in enumerate(merge_groups):
        print(f"    组{i+1}: 行 {group_start+1}-{group_end+1} (共 {group_end-group_start+1} 行), 内容=\"{cell_texts[group_start][:30]}\"")

    # 第四步：使用 XML vMerge 属性进行合并
    for group_start, group_end in merge_groups:
        if group_end > group_start:
            try:
                # 获取起始单元格（在数据行中的索引，需要加1因为跳过了表头）
                start_row_idx = group_start + 1
                start_cell = rows[start_row_idx].cells[col_idx]
                start_cell_element = start_cell._element

                # 保存起始单元格的文本内容
                merged_text = cell_texts[group_start]

                print(f"  合并组: 行{group_start+1}-{group_end+1} (共{group_end-group_start+1}行), 内容前30字符=\"{merged_text[:30]}\"")

                # 先设置 vMerge 属性（在清空内容之前）
                # 设置起始单元格的 vMerge 属性为 "restart"
                tc_pr = start_cell_element.get_or_add_tcPr()
                v_merge = OxmlElement('w:vMerge')
                v_merge.set(qn('w:val'), 'restart')
                tc_pr.append(v_merge)

                # 设置其他单元格的 vMerge 属性为 "continue"
                for row_offset in range(1, group_end - group_start + 1):
                    curr_row_idx = start_row_idx + row_offset
                    curr_cell = rows[curr_row_idx].cells[col_idx]
                    curr_cell_element = curr_cell._element

                    merge_tc_pr = curr_cell_element.get_or_add_tcPr()
                    continue_vmerge = OxmlElement('w:vMerge')
                    continue_vmerge.set(qn('w:val'), 'continue')
                    merge_tc_pr.append(continue_vmerge)

                    merged_count += 1

                # 然后清空并设置内容（在设置 vMerge 之后）
                # 清空所有要合并的单元格的内容（除了第一个）
                for row_offset in range(1, group_end - group_start + 1):
                    curr_row_idx = start_row_idx + row_offset
                    curr_cell = rows[curr_row_idx].cells[col_idx]
                    for paragraph in curr_cell.paragraphs:
                        paragraph.text = ""

                # 在起始单元格中设置内容（不清空，只确保有内容）
                if merged_text:
                    # 检查起始单元格是否已有内容
                    current_text = start_cell.text.strip()
                    if not current_text or current_text != merged_text:
                        # 清空并重新设置
                        for paragraph in start_cell.paragraphs:
                            paragraph.text = ""
                        # 确保至少有一个段落
                        if not start_cell.paragraphs:
                            start_cell.add_paragraph()
                        start_cell.paragraphs[0].text = merged_text

                print(f"    完成: vMerge=\"restart\" 设置于行{start_row_idx}")
            except Exception as e:
                print(f"  合并单元格组时出错（行 {group_start+1}-{group_end+1}）: {e}")
                import traceback
                traceback.print_exc()

    # 第五步：设置所有单元格的对齐方式（居中）
    for row in table.rows:
        if col_idx < len(row.cells):
            cell = row.cells[col_idx]
            try:
                # 垂直居中对齐
                cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                # 段落居中对齐
                for paragraph in cell.paragraphs:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            except Exception as e:
                print(f"  设置单元格对齐方式时出错: {e}")

    print(f"  第 {col_idx} 列物理合并完成，合并了 {merged_count} 个单元格")
    return merged_count

    # 设置所有单元格的对齐方式（居中）
    for row in table.rows:
        if col_idx < len(row.cells):
            cell = row.cells[col_idx]
            try:
                # 垂直居中对齐
                cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                # 段落居中对齐
                for paragraph in cell.paragraphs:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            except Exception as e:
                print(f"  设置单元格对齐方式时出错: {e}")

    print(f"  第 {col_idx} 列物理合并完成，合并了 {merged_count} 个单元格")
    return merged_count


def merge_other_tables_vertical_cells(doc, context):
    """
    处理范围三类别表格的纵向单元格合并（XML方法）

    仅处理范围三各类别的详细表格，不包含表1和表2。

    Args:
        doc: Word文档对象
        context: 数据上下文字典（用于判断哪些表格有数据）
    """
    # 范围三类别名称映射
    category_names = {
        1: "购买的商品和服务",
        2: "资本商品",
        3: "燃料和能源相关活动",
        4: "上游运输和配送",
        5: "运营中产生的废弃物",
        6: "员工商务旅行",
        7: "员工通勤",
        8: "上游租赁资产",
        9: "下游运输和配送",
        10: "销售产品的加工",
        11: "销售产品的使用",
        12: "寿命终结处理",
        13: "下游租赁资产",
        14: "特许经营",
        15: "投资"
    }

    total_merged = 0

    # 记录已处理的表格索引，避免重复处理表1和表2
    processed_tables = set()

    # 对有数据的范围三类别表格进行处理
    for cat_num in range(1, 16):
        detail_items = context.get(f'scope3_category{cat_num}', [])
        emission_value = context.get(f'scope_3_category_{cat_num}_emissions', 0)

        # 只处理有数据的类别
        if (detail_items and len(detail_items) > 0) or (emission_value and emission_value > 0):
            category_name = category_names.get(cat_num, "")

            # 查找对应的表格
            table_idx = find_table_by_content(doc, [category_name, f'类别{cat_num}'])

            # 跳过表1和表2（索引0和1）以及已处理的表格
            if table_idx is not None and table_idx < len(doc.tables) and table_idx not in processed_tables:
                # 跳过表1和表2
                if table_idx < 2:
                    print(f"  跳过范围三类别{cat_num}表格（{category_name}，表格索引：{table_idx}），这是主表格")
                    continue

                table = doc.tables[table_idx]
                print(f"  找到范围三类别{cat_num}表格（{category_name}，表格索引：{table_idx}）")
                processed_tables.add(table_idx)

                try:
                    merged = merge_vertical_cells(table, 0)
                    total_merged += merged
                except Exception as e:
                    print(f"  处理范围三类别{cat_num}表格时出错: {e}")
                    import traceback
                    traceback.print_exc()

    print(f"  范围三类别表格合并总计：合并了 {total_merged} 个单元格")

    # 步骤4：检查并插入缺失的类别12排放因子表
    print(f"\\n  [步骤4] 检查类别12排放因子表...")
    cat12_items = context.get('cat12_ef_items', [])

    if cat12_items:
        # 检查是否已有类别12排放因子表
        has_cat12_table = False
        cat12_table_idx = None

        for i, table in enumerate(doc.tables):
            for row in table.rows[:3]:
                for cell in row.cells:
                    if '类别12' in cell.text and '排放因子' in cell.text:
                        has_cat12_table = True
                        cat12_table_idx = i
                        break
                if has_cat12_table:
                    break
            if has_cat12_table:
                break

        if not has_cat12_table:
            print(f"  警告: 类别12有数据({len(cat12_items)}条)但无表格，尝试插入...")

            # 找到类别11的排放因子表（更宽松的搜索条件）
            cat11_table_idx = None
            for i, table in enumerate(doc.tables):
                # 检查表格是否有"外销产品使用"或"cat11"相关内容
                table_text = ''
                for row in table.rows[:5]:
                    for cell in row.cells:
                        table_text += cell.text + ' '

                if '外销产品使用' in table_text or 'cat11' in table_text.lower():
                    # 进一步确认这是排放因子表（有"Activity name"或"低位发热量"等列）
                    if 'Activity name' in table_text or '低位发热量' in table_text:
                        cat11_table_idx = i
                        break

            if cat11_table_idx is not None:
                print(f"  找到类别11表格: 索引{cat11_table_idx}，在其后插入类别12表格")
                # 在类别11表格后插入类别12表格
                insert_cat12_emission_factor_table(doc, cat11_table_idx, cat12_items)
            else:
                print(f"  错误: 未找到类别11排放因子表，无法确定插入位置")
        else:
            print(f"  类别12排放因子表已存在: 表格{cat12_table_idx}")
    else:
        print(f"  类别12无数据，跳过")


def merge_table_vertical_cells(doc, context):
    """
    处理文档中表格的纵向单元格合并

    针对文档中的"表1"和"表2"（即范围一、二、三的明细表），
    调用 merge_vertical_cells 处理第一列（类别列）。

    Args:
        doc: Word文档对象
        context: 数据上下文字典（用于判断哪些表格有数据）
    """
    # 定义要处理的表格标识关键词
    table_keywords = [
        # 表1：范围一直接排放源表格
        {'keywords': ['范围一', '直接', '排放源'], 'name': '表1（范围一）'},
        # 表2：范围二三间接排放源表格
        {'keywords': ['范围二', '范围三', '间接', '排放源'], 'name': '表2（范围二三）'}
    ]

    total_merged = 0

    for table_info in table_keywords:
        # 根据关键词查找表格
        table_idx = find_table_by_content(doc, table_info['keywords'])

        if table_idx is not None and table_idx < len(doc.tables):
            table = doc.tables[table_idx]
            print(f"  找到{table_info['name']}（表格索引：{table_idx}）")

            # 处理第一列（类别列）的纵向合并
            try:
                merged = merge_vertical_cells(table, 0)
                total_merged += merged
            except Exception as e:
                print(f"  处理{table_info['name']}时出错: {e}")
                import traceback
                traceback.print_exc()
        else:
            print(f"  未找到{table_info['name']}，跳过")

    # 处理范围三各类别的详细表格（如果有）
    # 范围三类别名称映射
    category_names = {
        1: "购买的商品和服务",
        2: "资本商品",
        3: "燃料和能源相关活动",
        4: "上游运输和配送",
        5: "运营中产生的废弃物",
        6: "员工商务旅行",
        7: "员工通勤",
        8: "上游租赁资产",
        9: "下游运输和配送",
        10: "销售产品的加工",
        11: "销售产品的使用",
        12: "寿命终结处理",
        13: "下游租赁资产",
        14: "特许经营",
        15: "投资"
    }

    # 对有数据的范围三类别表格进行处理
    for cat_num in range(1, 16):
        detail_items = context.get(f'scope3_category{cat_num}', [])
        emission_value = context.get(f'scope_3_category_{cat_num}_emissions', 0)

        # 只处理有数据的类别
        if (detail_items and len(detail_items) > 0) or (emission_value and emission_value > 0):
            category_name = category_names.get(cat_num, "")

            # 查找对应的表格
            table_idx = find_table_by_content(doc, [category_name, f'类别{cat_num}'])

            if table_idx is not None and table_idx < len(doc.tables):
                table = doc.tables[table_idx]
                print(f"  找到范围三类别{cat_num}表格（表格索引：{table_idx}）")

                # 处理第一列（类别列）的纵向合并
                try:
                    merged = merge_vertical_cells(table, 0)
                    total_merged += merged
                except Exception as e:
                    print(f"  处理范围三类别{cat_num}表格时出错: {e}")

    print(f"  表格纵向合并总计：合并了 {total_merged} 个单元格")


if __name__ == "__main__":
    import sys

    # 检查命令行参数
    if len(sys.argv) > 1 and sys.argv[1] == '--generate':
        # 生成报告模式
        xlsx_path = sys.argv[2] if len(sys.argv) > 2 else "test_data.xlsx"
        output_path = sys.argv[3] if len(sys.argv) > 3 else "carbon_report.docx"
        generate_report_from_xlsx(xlsx_path=xlsx_path, output_path=output_path)
    else:
        # 默认执行生成报告
        print("使用 'python main.py --generate' 生成报告")
        generate_report_from_xlsx()