# -*- coding: utf-8 -*-
"""
升级 template.docx 以适配新的 report_config.py 嵌套结构
"""
import sys
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH


def upgrade_template():
    print("开始升级 template.docx...")

    doc = Document('template.docx')

    # 找到并替换量化方法说明章节
    # 我们需要找到特定的段落并替换它们

    new_sections = []

    # 遍历所有段落，找到需要替换的部分
    i = 0
    while i < len(doc.paragraphs):
        para = doc.paragraphs[i]
        text = para.text.strip()

        # 检测是否是量化方法说明的开始
        if text == '量化方法说明':
            print(f"找到量化方法说明章节，位置: {i}")

            # 删除旧内容并添加新结构
            # 从这里开始到 {% endif %} 结束的整个段落块需要替换

            # 找到这个章节的结束位置（找到下一个同级标题或文档结尾）
            end_idx = i + 1
            bracket_count = 0
            in_if_block = False

            # 找到整个量化方法说明块的结束
            while end_idx < len(doc.paragraphs):
                t = doc.paragraphs[end_idx].text
                if '{%' in t and 'endif' in t:
                    # 检查是否是最外层的 endif
                    if not in_if_block:
                        break
                    in_if_block = False
                if '{%' in t and 'if' in t and 'flags' in t:
                    in_if_block = True
                end_idx += 1

            end_idx = min(end_idx + 1, len(doc.paragraphs))
            print(f"量化方法章节结束位置: {end_idx}")

            # 保存要添加的新内容
            new_content_index = i
            break
        i += 1

    # 现在我们需要创建新的量化方法说明内容
    # 由于直接修改 Word 文档比较复杂，我们采用不同的策略：
    # 找到并替换特定的 Jinja2 标签

    print("\\n开始替换量化方法标签...")

    # 替换策略：找到旧的循环标签并替换为新的结构化标签
    for para in doc.paragraphs:
        text = para.text

        # 替换范围一的循环
        if '{% for source_name, method in quantification_methods.scope_1.items() %}' in text:
            # 这是旧的循环开始
            # 我们需要替换整个块
            new_text = text.replace(
                '{% for source_name, method in quantification_methods.scope_1.items() %}',
                '{% for source_key, method_info in quantification_methods.scope_1.items() %}'
            )
            # 清除段落内容并添加新文本
            para.clear()
            para.add_run(new_text)

        # 替换简单的 source_name: method 为结构化输出
        if '{{ source_name }}：{{ method }}' in text:
            para.clear()
            # 添加新的结构化内容
            p = para
            run = p.add_run('排放源：{{ method_info.name }}')
            run.bold = True

            # 添加模型段落
            p2 = para.insert_paragraph_before('')
            p2.add_run('（1）量化模型：{{ method_info.model }}')

            # 添加活动数据段落
            p3 = para.insert_paragraph_before('')
            p3.add_run('（2）活动数据AD：{{ method_info.ad }}')

            # 添加排放因子段落
            p4 = para.insert_paragraph_before('')
            p4.add_run('（3）排放因子EF：{{ method_info.ef }}')

            # 删除原来的段落内容
            para._element.getparent().remove(para._element)

    # 类似地处理范围二
    for para in doc.paragraphs:
        text = para.text
        if '{% for source_name, method in quantification_methods.scope_2.items() %}' in text:
            para.clear()
            para.add_run(
                '{% for source_key, method_info in quantification_methods.scope_2.items() %}'
            )

    # 处理范围三的引用
    for para in doc.paragraphs:
        text = para.text
        if '{{ quantification_methods.scope_3[cat_key] }}' in text:
            para.clear()
            # 范围三的输出需要是完整的
            para.add_run('''排放源：{{ quantification_methods.scope_3[cat_key].name }}

（1）量化模型：{{ quantification_methods.scope_3[cat_key].model }}

（2）活动数据AD：{{ quantification_methods.scope_3[cat_key].ad }}

（3）排放因子EF：{{ quantification_methods.scope_3[cat_key].ef }}''')

    # 由于直接修改 Word 文档很复杂，我们使用更直接的方法
    # 读取文档内容，替换文本，然后写回

    print("模板升级完成！")
    print("注意：由于 Word 文档的复杂性，可能需要手动调整格式。")

    return doc


# 由于直接操作 Word 文档比较复杂，我们采用更简单的方法
# 直接替换 XML 中的文本内容
def simple_replace_template():
    """
    简单替换方法：直接替换文档中的文本
    """
    print("使用简单替换方法升级模板...")

    doc = Document('template.docx')

    # 需要替换的映射
    replacements = [
        # 范围一循环开始
        ('{% for source_name, method in quantification_methods.scope_1.items() %}',
         '{% for source_key, method_info in quantification_methods.scope_1.items() %}'),

        # 范围一简单输出 -> 精细化输出（这部分需要在段落级别处理）
        ('{{ source_name }}：{{ method }}',
         '排放源：{{ method_info.name }}\\n（1）量化模型：{{ method_info.model }}\\n（2）活动数据AD：{{ method_info.ad }}\\n（3）排放因子EF：{{ method_info.ef }}'),

        # 范围二
        ('{% for source_name, method in quantification_methods.scope_2.items() %}',
         '{% for source_key, method_info in quantification_methods.scope_2.items() %}'),

        # 范围三引用
        ('{{ quantification_methods.scope_3[cat_key] }}',
         '{{ quantification_methods.scope_3[cat_key].name }}\\n（1）量化模型：{{ quantification_methods.scope_3[cat_key].model }}\\n（2）活动数据AD：{{ quantification_methods.scope_3[cat_key].ad }}\\n（3）排放因子EF：{{ quantification_methods.scope_3[cat_key].ef }}'),
    ]

    # 对所有段落执行替换
    for para in doc.paragraphs:
        for old, new in replacements:
            if old in para.text:
                # 清除段落并添加新内容
                runs = list(para.runs)
                for run in runs:
                    run.text = run.text.replace(old, new)

    # 对所有表格中的单元格执行替换
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    for old, new in replacements:
                        if old in para.text:
                            runs = list(para.runs)
                            for run in runs:
                                run.text = run.text.replace(old, new)

    # 保存为新文件
    doc.save('template_new.docx')
    print("已创建 template_new.docx")

    return True


if __name__ == '__main__':
    try:
        simple_replace_template()
        print("\\n升级完成！")
        print("新文件: template_new.docx")
    except Exception as e:
        print(f"升级失败: {e}")
        import traceback
        traceback.print_exc()
